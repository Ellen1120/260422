"""
STM 문서 규칙 기반 파서 (API 불필요)
python-docx + 정규식으로 조제 정보·초자·시약량 추출
"""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document

# word namespace
_W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_W_T  = f'{{{_W_NS}}}t'
_W_TR = f'{{{_W_NS}}}tr'
_W_TC = f'{{{_W_NS}}}tc'

BASE_DIR = Path(__file__).resolve().parent
_NETWORK_AI = Path(r"\\file\04. 품질본부\3. 품질관리담당\1. 담당 공용\AI")
_NETWORK_STM = _NETWORK_AI / "STM"
_LOCAL_STM = BASE_DIR.parent / "STM"
STM_DIR = _NETWORK_STM if _NETWORK_STM.exists() else _LOCAL_STM
STM_FOLDER = STM_DIR
DATA_FOLDER = BASE_DIR / "data"
KB_PATH = DATA_FOLDER / "knowledge_base.json"

# ── 정규식 상수 ───────────────────────────────────────────

_RE_KOREAN = re.compile(r"[가-힣]")

# 시험항목 키워드: 전체 줄이 일치해야 함 (줄 끝 앵커 $)
_RE_TEST_ITEM = re.compile(
    r"^(?:"
    r"Description"
    r"|Identification(?:\s+by\s+HPLC)?\s*"
    r"|Dissolution"
    r"|Assay"
    r"|(?:Uniformity\s+of\s+dosage\s+units?|Content\s+uniformity)(?:\s*\([^)]*\))?"
    r"|Related\s+substances?(?:\s*\(Method\s+[AB]\))?"
    r"|Water\s+content\s+by\s+KF"
    r"|Polymorphism\s+by\s+PXRD"
    r"|Microbial\s+Enumeration\s+Test"
    r"|Specified\s+Microorganisms\s+Test"
    r")(?:\s*<[^>]+>)?\s*$",
    re.IGNORECASE,
)

# 한글 시험항목 이름 → 영문 정규 이름 매핑
_KO_TEST_ITEM_MAP: list[tuple[re.Pattern, str]] = [
    (re.compile(r'^성상\s*$'), 'Description'),
    (re.compile(r'^확인시험\s*(?:\(HPLC\))?\s*$'), 'Identification'),
    (re.compile(r'^용출\s*$'), 'Dissolution'),
    (re.compile(r'^함량\s*$'), 'Assay'),
    (re.compile(r'^제제\s*균일성\s*(?:\([^)]+\))?\s*$'), 'Uniformity of dosage units (Content Uniformity)'),
    (re.compile(r'^유연물질\s*(?:\(방법\s*[IVXivx가-힣]+\))?\s*$'), 'Related substances'),
    (re.compile(r'^수분\s*(?:함량|측정)?\s*(?:\(KF\))?\s*$'), 'Water content by KF'),
    (re.compile(r'^결정형\s*(?:\(PXRD\))?\s*$'), 'Polymorphism by PXRD'),
    (re.compile(r'^미생물\s*한도시험\s*$'), 'Microbial Enumeration Test'),
    (re.compile(r'^특정\s*미생물\s*시험\s*$'), 'Specified Microorganisms Test'),
]


def _normalize_ko_strength(line: str) -> str | None:
    """'26 밀리그램/5 밀리그램' → '26 mg/5 mg', 한글 함량이 아니면 None."""
    if not _RE_KO_STRENGTH_LINE.match(line):
        return None
    nums = re.findall(r'\d+(?:\.\d+)?', line)
    return "/".join(f"{n} mg" for n in nums)


def _get_canonical_test_item_name(line: str) -> str | None:
    """영문 또는 한글 시험항목 이름이면 정규 영문명 반환, 아니면 None."""
    stripped = line.strip()
    if _RE_TEST_ITEM.match(stripped) and len(stripped.split()) <= 12:
        # "<Method-A: ...>" suffix에서 Method 식별자(A/B) 추출 → 별도 섹션으로 분리
        method_m = re.search(r'<Method[-\s]*([AB])\b', stripped, re.IGNORECASE)
        normalized = re.sub(r'\s*<[^>]+>\s*$', '', stripped).strip()
        if re.match(r'^content\s+uniformity\s*$', normalized, re.IGNORECASE):
            return 'Uniformity of dosage units (Content Uniformity)'
        base = normalized if normalized else stripped
        if method_m:
            return f"{base} (Method {method_m.group(1).upper()})"
        return base
    for pat, name in _KO_TEST_ITEM_MAP:
        if pat.match(stripped):
            return name
    return None


# 조제 헤딩: "preparation" 단어를 포함하고 짧은 줄 (내용줄 제외)
_RE_PREP_CONTENT_START = re.compile(
    r"^(?:Accurately\s+)?(?:Weigh|Transfer|Add|Pipette|Dissolve|Use|Place|Adjust|Take|Shake|Incubate"
    r"|Mix|Dilute|Note[:\s]|Centrifuge|Filter|Inject"
    r"|Proceed|Compare|Observe|Streak|Perform|After|The\s|If\s|Put\s|Pour\s|Blank\s+(?!preparation\b)"
    r"|Withdraw|Heat|Cool|Rinse)",
    re.IGNORECASE,
)

# 함량 패턴 (standalone: "30 mg/25 mg", "50 mg")
_RE_STRENGTH_LINE = re.compile(
    r"^\s*(\d+(?:\.\d+)?\s*mg(?:\s*/\s*\d+(?:\.\d+)?\s*mg)*)\s*$",
    re.IGNORECASE,
)
# preparation (for X mg)
_RE_STRENGTH_PAREN = re.compile(
    r"\bpreparation\s*\(\s*for\s+(.+?)\s*\)",
    re.IGNORECASE,
)
# 한글 함량 표기: "26 밀리그램/5 밀리그램"
_RE_KO_STRENGTH_LINE = re.compile(
    r"^\s*\d+(?:\.\d+)?\s*밀리그램(?:\s*/\s*\d+(?:\.\d+)?\s*밀리그램)*\s*$"
)
# 한글 검액/표준액 조제 헤딩에서 함량 추출: "검액 조제 (10/20/10 mg)" 또는 "검액 조제 (5/20/10, 10/10/10, 5/5/10 mg)"
_RE_KO_PREP_STRENGTH_HEADING = re.compile(
    r'(?:검액|표준액)\s*조제\s*\(([^)]+mg[^)]*)\)',
    re.IGNORECASE,
)
# 한글 용출 방법 헤딩: "방법 I: 에제티미브", "방법 II: 로수바스타틴"
_RE_KO_METHOD_HEADING = re.compile(
    r'^방법\s+([IVX]+)\s*(?::\s*(.+))?\s*$',
    re.IGNORECASE,
)

# 시약 추출 패턴 (순서 중요: 구체적인 것 먼저)
_INGREDIENT_PATTERNS: list[tuple[str, re.Pattern]] = [
    # Weigh/Transfer X mg/g of [name] into/and/,
    ("weigh_transfer",
     re.compile(
         r"(?:Accurately\s+)?(?:Weigh\s+and\s+transfer|Transfer|Dissolve|Weigh)"
         r"\s+(?:about\s+)?(\d+(?:\.\d+)?)\s*(mg|g)\s+of\s+"
         r"([A-Za-z0-9][^\n]{2,60}?)"
         r"(?:\s+(?:standard\b|into\b|to\b|in\b|and\b)|[,\.])",
         re.IGNORECASE,
     )),
    # X mL of [name] in Y mL
    ("in_ml",
     re.compile(
         r"(\d+(?:\.\d+)?)\s*(mL|L)\s+of\s+"
         r"([A-Za-z0-9][^\n]{2,60}?)"
         r"\s+in\s+\d",
         re.IGNORECASE,
     )),
    # Pipette X mL of [name]
    ("pipette",
     re.compile(
         r"Pipette\s+(\d+(?:\.\d+)?)\s*(mL|L)\s+of\s+"
         r"([A-Za-z0-9][^\n]{2,60}?)"
         r"(?:\s+(?:into\b|to\b|and\b)|[,\.])",
         re.IGNORECASE,
     )),
    # Add about X mL/mg/g of [name]
    ("add",
     re.compile(
         r"Add\s+(?:about\s+)?(\d+(?:\.\d+)?)\s*(mL|mg|g|L)\s+of\s+"
         r"([A-Za-z0-9][^\n]{2,60}?)"
         r"(?:\s+(?:and\b|to\b|,|\.))",
         re.IGNORECASE,
     )),
    # Dilute X mL of [substance] with water to Y mL  (희석)
    ("dilute",
     re.compile(
         r"Dilute\s+(\d+(?:\.\d+)?)\s*(mL|L)\s+of\s+"
         r"([A-Za-z0-9][^\n]{2,40}?)"
         r"\s+with\s+\w+\s+to\s+\d",
         re.IGNORECASE,
     )),
    # Dissolve X g/mg of [substance] in Y mL/L of [solvent]
    ("dissolve_in",
     re.compile(
         r"Dissolve\s+(?:about\s+)?(\d+(?:\.\d+)?)\s*(mg|g)\s+of\s+"
         r"([A-Za-z0-9][^\n]{2,60}?)\s+in\s+(?:about\s+)?(?:\d+(?:\.\d+)?\s*(?:mL|L)\s+of\s+)?([A-Za-z0-9][^\n]{2,60}?)",
         re.IGNORECASE,
     )),
]

# Mix A and B in the ratio of X:Y v/v
_RE_RATIO = re.compile(
    r"Mix\s+(.+?)\s+and\s+(.+?)\s+in\s+the\s+ratio\s+of\s+"
    r"(\d+(?:\.\d+)?)\s*:\s*(\d+(?:\.\d+)?)",
    re.IGNORECASE,
)

# 한글 시약 추출 패턴
_KO_INGREDIENT_PATTERNS: list[re.Pattern] = [
    # 일반: 한글/영문으로 시작하는 이름 (mg/g)
    re.compile(
        r'([가-힣A-Za-z][가-힣A-Za-z0-9\s\-()]+?)\s+(?:약\s*)?(\d+(?:\.\d+)?)\s*(mg|g)\s*(?:[을를과]|달아)',
        re.IGNORECASE,
    ),
    # "1-헥산설폰산나트륨" 등 숫자-하이픈으로 시작하는 IUPAC 스타일 이름 (mg/g)
    re.compile(
        r'(\d+\-[가-힣A-Za-z][가-힣A-Za-z0-9\s\-()]*?)\s+(?:약\s*)?(\d+(?:\.\d+)?)\s*(mg|g)\s*(?:[을를과]|달아)',
        re.IGNORECASE,
    ),
    # 일반: 한글/영문으로 시작하는 이름 (mL) — "에" 포함 (e.g. "아세트산 2.9 mL에")
    re.compile(
        r'([가-힣A-Za-z][가-힣A-Za-z0-9\s\-()]+?)\s+(\d+(?:\.\d+)?)\s*(mL)\s*(?:[을를]|및|까지|에)',
        re.IGNORECASE,
    ),
]

# 한글 비율 혼합 패턴 (2성분: "A와 B를 각각 X:Y")
_RE_RATIO_KO2 = re.compile(
    r'([가-힣A-Za-z][가-힣A-Za-z0-9\s\-()]*?)\s*(?:와|과)\s*'
    r'([가-힣A-Za-z][가-힣A-Za-z0-9\s\-()]*?)\s*(?:을|를)?\s*각각\s*'
    r'(\d+(?:\.\d+)?)\s*:\s*(\d+(?:\.\d+)?)',
)
# 한글 3성분 비율: "A, B 및 C를 X:Y:Z v/v/v"
_RE_RATIO_KO3 = re.compile(
    r'([가-힣A-Za-z][가-힣A-Za-z0-9\s\-()]*?)\s*,\s*'
    r'([가-힣A-Za-z][가-힣A-Za-z0-9\s\-()]*?)\s+및\s+'
    r'([가-힣A-Za-z][가-힣A-Za-z0-9\s\-()]*?)\s*(?:을|를)?\s*'
    r'(\d+(?:\.\d+)?)\s*:\s*(\d+(?:\.\d+)?)\s*:\s*(\d+(?:\.\d+)?)',
)
# 한글 2성분 비율 (및, 각각 없음): "A 및 B를 X:Y v/v"
_RE_RATIO_KO2_MIT = re.compile(
    r'([가-힣A-Za-z][가-힣A-Za-z0-9\s\-()]*?)\s+및\s+'
    r'([가-힣A-Za-z][가-힣A-Za-z0-9\s\-()]*?)\s*(?:을|를)?\s*'
    r'(\d+(?:\.\d+)?)\s*:\s*(\d+(?:\.\d+)?)'
    r'(?!\s*:\s*\d)',
)
# 한글 2성분 비율 (와/과 + 비율로 혼합): "A와 B를 X:Y (v/v)의 비율로 혼합"
_RE_RATIO_KO2_WA = re.compile(
    r'([가-힣A-Za-z][가-힣A-Za-z0-9\s\-()]*?)\s*(?:와|과)\s*'
    r'([가-힣A-Za-z][가-힣A-Za-z0-9\s\-()]*?)\s*(?:을|를)?\s*'
    r'(\d+(?:\.\d+)?)\s*:\s*(\d+(?:\.\d+)?)'
    r'(?!\s*:\s*\d)'
    r'(?:\s*\([vV]/[vV]\))?\s*(?:의)?\s*비율로',
)
# 콜론 구분 3성분 비율: "A:B:C = X:Y:Z (v/v/v)" 또는 "A:B:C X:Y:Z"
_RE_RATIO_COLON3 = re.compile(
    r'([A-Za-z가-힣][A-Za-z가-힣0-9\s\-()]*?)\s*:\s*'
    r'([A-Za-z가-힣][A-Za-z가-힣0-9\s\-()]*?)\s*:\s*'
    r'([A-Za-z가-힣][A-Za-z가-힣0-9\s\-()]*?)\s*'
    r'(?:=\s*)?(\d+(?:\.\d+)?)\s*:\s*(\d+(?:\.\d+)?)\s*:\s*(\d+(?:\.\d+)?)',
)

# 초자 패턴 (영문)
_RE_GLASSWARE = re.compile(
    r"(\d+(?:\.\d+)?)\s*mL\s+(?:of\s+)?(?:a\s+)?"
    r"(volumetric\s+flask|graduated\s+cylinder|measuring\s+cylinder"
    r"|pipette|burette|beaker|vial|용량플라스크|메스플라스크|피펫|메스실린더)",
    re.IGNORECASE,
)

# [동사] + [용량] 패턴 (Pipette 1 mL, Transfer 5 mL 등)
_RE_VERB_GLASSWARE = re.compile(
    r"(?:Pipette|Transfer|Take)\s+(\d+(?:\.\d+)?)\s*(?:mL|L)",
    re.IGNORECASE
)

# 초자 패턴 (한글) — 차광 수식어 및 L 단위, "용량 플라스크"(띄어쓰기) 허용
_RE_GLASSWARE_KO = re.compile(
    r"(\d+(?:\.\d+)?)\s*(mL|L)\s*(?:차광\s*)?(용량\s*플라스크|메스실린더|메스플라스크|피펫|비이커|바이알)",
)
# 한글 전달 피펫: "10 mL를 정확하게 취하여"
_RE_KO_PIPETTE = re.compile(
    r"(\d+(?:\.\d+)?)\s*mL\s*를\s*(?:정확하게|정밀하게)?\s*취하여",
)

# 원심분리 팔콘 검출
_RE_CENTRIFUGE = re.compile(r"\bcentrifuge\b", re.IGNORECASE)

# 시린지 필터: 재질 명시형 (0.45 µm PVDF (Millipore))
_RE_FILTER_MENTION = re.compile(
    r"(\d+(?:\.\d+)?)\s*[uμµ]\s*m\s+(PVDF|Nylon|PTFE|PES|MCE|Cellulose)\s*(?:\(([^)]+)\))?",
    re.IGNORECASE,
)
# 시린지 필터: 재질 미명시형 (0.45µm syringe filter)
_RE_GENERIC_SYRINGE_FILTER = re.compile(
    r"(\d+(?:\.\d+)?)\s*[uμµ]\s*m\s+syringe\s+filter",
    re.IGNORECASE,
)
# 시린지 필터: 한글 표현 (0.45 µm PTFE 시린지 필터)
_RE_FILTER_KO = re.compile(
    r"(\d+(?:\.\d+)?)\s*[uμµ]\s*m\s+(PVDF|Nylon|PTFE|PES|MCE)?\s*시린지\s*필터",
    re.IGNORECASE,
)

# 용출 표준액 조제에서 volumetric flask 크기 추출 (영문/한글)
_RE_VOLUMETRIC_FLASK_SIZE = re.compile(
    r"(\d+(?:\.\d+)?)\s*mL\s+(?:of\s+)?volumetric\s+flask",
    re.IGNORECASE,
)
_RE_KO_VOLUMETRIC_FLASK_SIZE = re.compile(
    r"(\d+(?:\.\d+)?)\s*mL\s+(?:용량플라스크|메스플라스크)",
)

# 최종 볼륨 추출 (우선순위 순)
_RE_VOLUME_PATTERNS: list[re.Pattern] = [
    re.compile(r"dilute\s+to\s+(\d+(?:\.\d+)?)\s*mL", re.IGNORECASE),
    re.compile(r"with\s+\w[\w\s-]*?\s+to\s+(\d+(?:\.\d+)?)\s*mL", re.IGNORECASE),
    re.compile(r"into\s+(?:a\s+)?(\d+(?:\.\d+)?)\s*mL\s+(?:of\s+)?volumetric\s+flask", re.IGNORECASE),
    re.compile(r"(\d+(?:\.\d+)?)\s*mL\s+volumetric\s+flask", re.IGNORECASE),
    re.compile(r"into\s+(\d+(?:\.\d+)?)\s*mL\s+(?:of\s+)?(?:Milli-Q[-\s]?water|purified\s+water)", re.IGNORECASE),
    re.compile(r"in\s+(\d+(?:\.\d+)?)\s*mL\s+of\s+(?:purified|Milli-Q|water)", re.IGNORECASE),
    re.compile(r"in\s+(\d+(?:\.\d+)?)\s*mL\s+of\s+\w", re.IGNORECASE),
    # 한글: "X mL 용량플라스크에 넣고 표선" or "물 X mL에 넣고"
    re.compile(r"(\d+(?:\.\d+)?)\s*mL\s+용량플라스크"),
    re.compile(r"물\s+(\d+(?:\.\d+)?)\s*mL에\s+(?:넣|녹|옮)"),
    # 한글: "X mL 차광 용량플라스크" / "X mL차광용량플라스크" (차광 수식어 포함)
    re.compile(r"(\d+(?:\.\d+)?)\s*mL\s*차광\s*(?:용량\s*플라스크|메스플라스크)"),
]

# 제품명 추출
_RE_PRODUCT_NAME_WS = [
    re.compile(r"Ws\s*:.*?as\s+([A-Za-z][A-Za-z0-9\s\-]+?)(?:\s+sodium\b|\s+hydrate\b|\s+on\s+as|\s+\(|[,\.]|$)", re.IGNORECASE),
    re.compile(r"Ws\s*:\s*Weight\s+of\s+([A-Za-z][A-Za-z0-9\s\-]+?)\s+standard\s+taken", re.IGNORECASE),
    re.compile(r"Weigh\s+and\s+transfer\s+(?:about\s+)?[\d\.]+\s*(?:mg|g)\s+of\s+([A-Za-z][A-Za-z0-9\s\-]+?)\s+standard", re.IGNORECASE),
]


# ── 유틸 함수 ─────────────────────────────────────────────

def _is_korean(text: str) -> bool:
    return bool(_RE_KOREAN.search(text))


def _is_reagent_solution(line: str) -> bool:
    """조제명이 시약(Diluent, MP 등)인지 여부 판별."""
    line_lower = line.lower()
    # 시약 키워드
    REAGENT_KEYWORDS = ["diluent", "mobile phase", "buffer", "희석액", "이동상", "완충액", "용액", "시약"]
    # 시험 용액 키워드 (이게 포함되면 시약이 아님)
    SAMPLE_KEYWORDS = ["standard", "sample", "test", "stock", "internal", "linearity", "quantitation", "resolution", "표준", "검액", "원액", "내부표준", "대조"]
    
    if any(kw in line_lower for kw in SAMPLE_KEYWORDS):
        return False
    return any(kw in line_lower for kw in REAGENT_KEYWORDS)


def _is_prep_heading(line: str) -> bool:
    """조제 섹션 헤딩 여부 판별. 이제 모든 조제를 포함하되 구분만 함."""
    line_lower = line.lower().strip()

    # 한글: "XX 조제" 또는 "XX 조제 (함량/조건)" 패턴 — 점·%·+ 등 특수문자 포함 가능
    if (re.search(r'조제\s*(?:\([^)]*\))?\s*$', line) and
            len(line) <= 80 and
            not _RE_PREP_CONTENT_START.match(line)):
        return True

    # 영문: "preparation" 단어가 포함된 헤딩
    if "preparation" in line_lower:
        if _RE_PREP_CONTENT_START.match(line):
            return False
        if line_lower.startswith("note"):
            return False
        if len(line) > 80:
            return False
        return True

    return False


def _derive_solution_name(heading: str) -> str:
    """'Buffer preparation' → 'Buffer', '완충액 조제' → '완충액', '검액 조제 (10/20/10 mg)' → '검액 (10/20/10 mg)'"""
    if '조제' in heading:
        # "BASE 조제" or "BASE 조제 (PAREN)" format
        m = re.match(r'^(.+?)\s+조제(?:\s*(\([^)]*\)))?\s*$', heading)
        if m:
            base = m.group(1).strip()
            paren = m.group(2) or ''
            return (base + ' ' + paren).strip() if paren else base
        return re.sub(r'\s*조제\s*', ' ', heading).strip()
    # 영문 "preparation" 제거
    name = re.sub(r"\s+preparation\b", "", heading, flags=re.IGNORECASE).strip()
    return name.strip()


# ── 핵심 추출 함수 ────────────────────────────────────────

def _extract_volume(heading: str, lines: list[str]) -> float | None:
    """준비 텍스트에서 최종 조제 볼륨(mL) 추출. 명시적 볼륨 없으면 비율 합산으로 대체."""
    combined = heading + " " + " ".join(lines)
    for pat in _RE_VOLUME_PATTERNS:
        m = pat.search(combined)
        if m:
            return float(m.group(1))
    # 한글: "X L 용량플라스크" → X * 1000 mL
    m_l = re.search(
        r"(\d+(?:\.\d+)?)\s*L\s*(?:차광\s*)?(?:용량\s*플라스크|메스플라스크)",
        combined,
    )
    if m_l:
        return float(m_l.group(1)) * 1000
    # 비율 혼합이고 총량 미지정인 경우 비율 합산을 기본 볼륨으로 사용
    m = _RE_RATIO.search(combined)
    if m:
        return float(m.group(3)) + float(m.group(4))
    m3 = _RE_RATIO_KO3.search(combined)
    if m3:
        return float(m3.group(4)) + float(m3.group(5)) + float(m3.group(6))
    m2 = _RE_RATIO_KO2.search(combined) or _RE_RATIO_KO2_MIT.search(combined)
    if m2:
        return float(m2.group(3)) + float(m2.group(4))
    return None


def _extract_ingredients(text: str, final_volume_ml: float | None) -> list[dict]:
    """준비 절차 텍스트에서 시약 목록 추출."""
    results: list[dict] = []
    seen: set[tuple] = set()

    def _add(name: str, amount: float, unit: str) -> None:
        name = name.strip().rstrip(",").strip()
        # 한글 조사(을/를/이/가/은/는/도/와/과) 후행 제거
        name = re.sub(r'[을를이가은는도와과]\s*$', '', name).strip()
        # 절차 설명 제거: ", sonicate for..." / ". sonicate..."
        name = re.split(r"[,\.]\s+(?:sonicate|mix|stir|shake|heat|cool|filter)", name, flags=re.IGNORECASE)[0].strip()
        # "subsequent filtrate" 등 시약이 아닌 단어 제외
        if re.search(r"filtrate|residue|supernatant|layer|solution\s+from", name, re.IGNORECASE):
            return
        # "standard" 접미사 제거
        name = re.sub(r"\s+standard$", "", name, flags=re.IGNORECASE).strip()
        # 앞에 붙은 "the " 제거
        name = re.sub(r"^the\s+", "", name, flags=re.IGNORECASE).strip()
        if len(name) < 2 or len(name) > 70:
            return
        # 모호한 참조 건너뜀 (영문: "the above" 등, 한글: "이 액", "이 용액")
        if re.match(r"^(the\s+above|above|it|them|each|following|이\s+액|이\s+용액)$", name, re.IGNORECASE):
            return
        # 이름 내 단위 표기(mL, mg) 포함 → 절차 텍스트 오인식, 제외 (case-sensitive)
        if re.search(r'mL|mg', name):
            return
        key = (name.lower(), unit.lower())
        if key not in seen:
            seen.add(key)
            results.append({"name": name, "amount": round(amount, 4), "unit": unit.lower()})

    # 패턴 기반 추출 (영문)
    for _tag, pat in _INGREDIENT_PATTERNS:
        for m in pat.finditer(text):
            _add(m.group(3), float(m.group(1)), m.group(2))

    # 한글 시약 추출
    for pat in _KO_INGREDIENT_PATTERNS:
        for m in pat.finditer(text):
            name_raw = m.group(1).strip().rstrip(',').strip()
            _add(name_raw, float(m.group(2)), m.group(3))

    # 비율 혼합 영문 (Mix A and B in ratio X:Y)
    for m in _RE_RATIO.finditer(text):
        sub_a = m.group(1).strip()
        sub_b = m.group(2).strip()
        ratio_a = float(m.group(3))
        ratio_b = float(m.group(4))
        total = ratio_a + ratio_b
        if final_volume_ml and total > 0:
            _add(sub_a, round(ratio_a / total * final_volume_ml, 1), "mL")
            _add(sub_b, round(ratio_b / total * final_volume_ml, 1), "mL")

    # 비율 혼합 한글 2성분 (A와 B를 각각 X:Y)
    for m in _RE_RATIO_KO2.finditer(text):
        sub_a = m.group(1).strip()
        sub_b = m.group(2).strip()
        ratio_a = float(m.group(3))
        ratio_b = float(m.group(4))
        total = ratio_a + ratio_b
        if final_volume_ml and total > 0:
            _add(sub_a, round(ratio_a / total * final_volume_ml, 1), "mL")
            _add(sub_b, round(ratio_b / total * final_volume_ml, 1), "mL")

    # 비율 혼합 한글 3성분 (A, B 및 C를 X:Y:Z v/v/v)
    for m in _RE_RATIO_KO3.finditer(text):
        sub_a = m.group(1).strip()
        sub_b = m.group(2).strip()
        sub_c = m.group(3).strip()
        ratio_a = float(m.group(4))
        ratio_b = float(m.group(5))
        ratio_c = float(m.group(6))
        total = ratio_a + ratio_b + ratio_c
        if final_volume_ml and total > 0:
            _add(sub_a, round(ratio_a / total * final_volume_ml, 1), "mL")
            _add(sub_b, round(ratio_b / total * final_volume_ml, 1), "mL")
            _add(sub_c, round(ratio_c / total * final_volume_ml, 1), "mL")

    # 비율 혼합 한글 2성분 (A 및 B를 X:Y, 각각 없음) - 3성분 텍스트 제외
    if not _RE_RATIO_KO3.search(text):
        for m in _RE_RATIO_KO2_MIT.finditer(text):
            sub_a = m.group(1).strip()
            sub_b = m.group(2).strip()
            ratio_a = float(m.group(3))
            ratio_b = float(m.group(4))
            total = ratio_a + ratio_b
            if final_volume_ml and total > 0:
                _add(sub_a, round(ratio_a / total * final_volume_ml, 1), "mL")
                _add(sub_b, round(ratio_b / total * final_volume_ml, 1), "mL")

    # 비율 혼합 한글 2성분 (A와 B를 X:Y v/v 비율로 혼합)
    for m in _RE_RATIO_KO2_WA.finditer(text):
        sub_a = m.group(1).strip()
        sub_b = m.group(2).strip()
        ratio_a = float(m.group(3))
        ratio_b = float(m.group(4))
        total = ratio_a + ratio_b
        if final_volume_ml and total > 0:
            _add(sub_a, round(ratio_a / total * final_volume_ml, 1), "mL")
            _add(sub_b, round(ratio_b / total * final_volume_ml, 1), "mL")

    # 콜론 구분 3성분: "A:B:C = X:Y:Z (v/v/v)"
    for m in _RE_RATIO_COLON3.finditer(text):
        sub_a, sub_b, sub_c = m.group(1).strip(), m.group(2).strip(), m.group(3).strip()
        ratio_a, ratio_b, ratio_c = float(m.group(4)), float(m.group(5)), float(m.group(6))
        total = ratio_a + ratio_b + ratio_c
        if final_volume_ml and total > 0:
            _add(sub_a, round(ratio_a / total * final_volume_ml, 1), "mL")
            _add(sub_b, round(ratio_b / total * final_volume_ml, 1), "mL")
            _add(sub_c, round(ratio_c / total * final_volume_ml, 1), "mL")

    return results


_KO_GLASSWARE_MAP = {
    '용량플라스크': 'volumetric flask',
    '메스실린더': 'graduated cylinder',
    '메스플라스크': 'volumetric flask',
    '피펫': 'pipette',
    '홀 피펫': 'pipette',
    '홀피펫': 'pipette',
    '비이커': 'beaker',
    '바이알': 'vial',
    '유발': 'mortar',
}

_RE_MORTAR = re.compile(r"\b(?:mortar|mortar\s+and\s+pestle)\b|유발", re.IGNORECASE)

def _normalize_glassware_size(s: str) -> str:
    """초자 규격 정규화 (예: 10.0 -> 10)"""
    try:
        val = float(s)
        if val == int(val):
            return str(int(val))
        return str(val)
    except:
        return s.strip()

def _extract_glassware(text: str) -> list[dict]:
    """준비 절차 텍스트에서 초자 목록 추출. 동일 규격이 여러 번 나오면 count_per_batch 합산."""
    count_map: dict[tuple, int] = {}
    
    # 1. 영문 패턴 추출
    for m in _RE_GLASSWARE.finditer(text):
        size = _normalize_glassware_size(m.group(1))
        gtype = m.group(2).lower()
        # 정규화
        if any(kw in gtype for kw in ["flask", "플라스크"]):
            gtype = "volumetric flask"
        elif any(kw in gtype for kw in ["cylinder", "실린더"]):
            gtype = "graduated cylinder"
        elif any(kw in gtype for kw in ["pipette", "피펫"]):
            gtype = "pipette"
        
        gtype = _KO_GLASSWARE_MAP.get(gtype, gtype)
        key = (gtype, size)
        count_map[key] = count_map.get(key, 0) + 1
    
    # 2. Pipette 동사형 패턴
    for m in _RE_VERB_GLASSWARE.finditer(text):
        size = _normalize_glassware_size(m.group(1))
        key = ("pipette", size)
        count_map[key] = count_map.get(key, 0) + 1

    # 3. 한글 패턴 추출 (차광 수식어, L 단위, 용량 플라스크 띄어쓰기 허용)
    for m in _RE_GLASSWARE_KO.finditer(text):
        raw_size = float(m.group(1))
        unit = m.group(2).lower()
        if unit == 'l':
            raw_size = raw_size * 1000  # L → mL 변환
        size = _normalize_glassware_size(str(int(raw_size) if raw_size == int(raw_size) else raw_size))
        gtype_raw = re.sub(r'\s+', '', m.group(3))  # "용량 플라스크" → "용량플라스크"
        gtype = _KO_GLASSWARE_MAP.get(gtype_raw, gtype_raw.lower())
        if any(kw in gtype for kw in ["flask", "플라스크"]):
            gtype = "volumetric flask"
        elif any(kw in gtype for kw in ["cylinder", "실린더"]):
            gtype = "graduated cylinder"
        elif any(kw in gtype for kw in ["pipette", "피펫"]):
            gtype = "pipette"

        key = (gtype, size)
        count_map[key] = count_map.get(key, 0) + 1

    # 4. 한글 취하여 피펫 패턴: "10 mL를 정확하게 취하여"
    for m in _RE_KO_PIPETTE.finditer(text):
        size = _normalize_glassware_size(m.group(1))
        key = ("pipette", size)
        count_map[key] = count_map.get(key, 0) + 1

    result = [
        {"type": gtype, "size": f"{size} mL", "count_per_batch": cnt}
        for (gtype, size), cnt in count_map.items()
    ]
    if _RE_MORTAR.search(text):
        result.append({"type": "mortar", "size": "-", "count_per_batch": 1})
    return result


def _extract_filters_from_text(text: str) -> list[dict]:
    """조제 텍스트에서 시린지 필터 추출 - 줄별로 처리해 복수 종류를 모두 포함."""
    filter_lines = [
        line for line in text.splitlines()
        if re.search(r"filter\s+through|syringe\s+filter|시린지\s*필터", line, re.IGNORECASE)
    ]
    if not filter_lines:
        return []

    results: list[dict] = []
    seen: set[tuple] = set()

    for line in filter_lines:
        specific_found = False
        for m in _RE_FILTER_MENTION.finditer(line):
            size = float(m.group(1))
            material = m.group(2).upper()
            mfr = m.group(3).strip() if m.group(3) else ""
            key = (size, material, mfr)
            if key not in seen:
                seen.add(key)
                results.append({
                    "size_um": size, "material": material,
                    "manufacturer": mfr, "filter_type": "syringe", "count_per_batch": 1,
                })
            specific_found = True
        if not specific_found:
            for m in _RE_GENERIC_SYRINGE_FILTER.finditer(line):
                size = float(m.group(1))
                key = (size, "", "")
                if key not in seen:
                    seen.add(key)
                    results.append({
                        "size_um": size, "material": "", "manufacturer": "",
                        "filter_type": "syringe", "count_per_batch": 1,
                    })
            for m in _RE_FILTER_KO.finditer(line):
                size = float(m.group(1))
                material = (m.group(2) or "").upper()
                key = (size, material, "")
                if key not in seen:
                    seen.add(key)
                    results.append({
                        "size_um": size, "material": material, "manufacturer": "",
                        "filter_type": "syringe", "count_per_batch": 1,
                    })
    return results


def _extract_product_names(paragraphs: list[str]) -> list[str]:
    """문서에서 제품명 추출 (복수 성분 지원)."""
    names: list[str] = []
    seen_lower: set[str] = set()

    for line in paragraphs:
        for pat in _RE_PRODUCT_NAME_WS:
            m = pat.search(line)
            if m:
                name = m.group(1).strip().rstrip(".,")
                # 숫자만 이거나 너무 짧으면 제외
                if re.match(r"^[\d\.\s]+$", name) or len(name) < 4:
                    continue
                # 불순물 표준품 제외
                if re.search(r"\bimpurit", name, re.IGNORECASE):
                    continue
                # 너무 긴 이름 제외 (변수 설명 등)
                if len(name) > 50:
                    continue
                # 제목화 (모두 소문자인 경우만)
                if name == name.lower():
                    name = name.title()
                if name.lower() not in seen_lower:
                    seen_lower.add(name.lower())
                    names.append(name)
    return names


def _extract_strengths(paragraphs: list[str]) -> list[str]:
    """문서에서 함량/규격 추출."""
    strengths: list[str] = []
    seen: set[str] = set()

    def _add(s: str) -> None:
        s = s.strip()
        if s and s not in seen:
            seen.add(s)
            strengths.append(s)

    for line in paragraphs:
        # "50 mg" or "30 mg/25 mg"
        m = _RE_STRENGTH_LINE.match(line)
        if m:
            _add(m.group(1).replace(" ", "").replace("mg", " mg").strip())
        # "preparation (for 50 mg)" — 복수 함량 "(For 80/10 mg, 40/10 mg)" 분리
        m2 = _RE_STRENGTH_PAREN.search(line)
        if m2:
            raw = m2.group(1).strip()
            parts = re.split(r'\s*,\s*|\s+&\s+|\s+and\s+', raw, flags=re.IGNORECASE)
            added = False
            for part in parts:
                part = part.strip()
                # "30 mg/25 mg" 또는 "80/10 mg" 두 가지 형식 모두 허용
                if (_RE_STRENGTH_LINE.match(part) or
                        re.match(r'^\d+(?:\.\d+)?(?:/\d+(?:\.\d+)?)+\s*mg\s*$', part, re.IGNORECASE)):
                    _add(part)
                    added = True
            if not added:
                _add(raw)

    # 한글 함량: "26 밀리그램/5 밀리그램"
    for line in paragraphs:
        ko_str = _normalize_ko_strength(line)
        if ko_str:
            _add(ko_str)

    # 한글 검액/표준액 조제 헤딩: "검액 조제 (10/20/10 mg)" 또는 "검액 조제 (5/20/10, 10/10/10, ..., 5/5/10 mg)"
    for line in paragraphs:
        m_kp = _RE_KO_PREP_STRENGTH_HEADING.search(line)
        if m_kp:
            raw = m_kp.group(1).strip()
            parts = re.split(r'\s*,\s*', raw)
            has_unit = bool(re.search(r'\bmg\b', parts[-1], re.IGNORECASE))
            for part in parts:
                part = part.strip()
                if has_unit and not re.search(r'\bmg\b', part, re.IGNORECASE):
                    part = part + ' mg'
                # "10mg" → "10 mg", "25/10mg" → "25/10 mg"
                part = re.sub(r'(\d)(mg)', r'\1 \2', part, flags=re.IGNORECASE).strip()
                if re.search(r'\d+', part):
                    _add(part)

    # 단일 함량 문서: "Standard solution preparation (for 50 mg)" 패턴 없을 때
    # Label claim 수치로부터 추출 시도 (예: "20 mg Vonoprazan" 또는 section header)
    if not strengths:
        for line in paragraphs:
            m3 = re.search(
                r"(?:for|of)\s+(\d+(?:\.\d+)?\s*mg(?:\s*/\s*\d+(?:\.\d+)?\s*mg)*)\s+"
                r"(?:tablet|capsule|strength|dose)",
                line, re.IGNORECASE
            )
            if m3:
                _add(m3.group(1).strip())

    return strengths or ["N/A"]


# ── 섹션 파싱 ────────────────────────────────────────────

def _parent_section_name(name: str, current_name: str | None) -> str | None:
    """
    name이 current_name의 하위 항목이면 current_name 반환, 아니면 None.
    예) 'Identification by HPLC' → current 'Identification' → 'Identification'
    """
    if current_name and name.lower().startswith(current_name.lower() + " "):
        return current_name
    return None


def _parse_sections(english_lines: list[str]) -> list[dict]:
    """
    영문 줄 목록을 시험항목 섹션으로 그룹화.
    하위 항목(e.g. 'Identification by HPLC')은 상위 항목('Identification')에 병합.
    반환: [{name, lines}]
    """
    sections: list[dict] = []
    current: dict | None = None

    for line in english_lines:
        canonical = _get_canonical_test_item_name(line)
        if canonical:
            parent = _parent_section_name(canonical, current["name"] if current else None)
            # 같은 정규명 또는 하위 항목이면 현재 섹션에 병합
            if parent or (current and canonical == current["name"]):
                current["lines"].append(line)
            else:
                if current:
                    sections.append(current)
                current = {"name": canonical, "lines": []}
        elif current is not None:
            current["lines"].append(line)

    if current:
        sections.append(current)
    return sections


def _is_any_prep_heading(line: str) -> bool:
    """모든 종류의 조제 헤딩인지 확인."""
    return _is_prep_heading(line)




def _parse_preparations(section_lines: list[str]) -> list[dict]:
    """
    섹션 내용 줄에서 조제 블록 추출.
    반환: [{..., is_reagent: bool}]
    """
    blocks: list[dict] = []
    current_heading: str | None = None
    current_lines: list[str] = []


    # CP029 용출 시약 오인식 방지를 위한 필터링 단어
    EXCLUDE_KEYWORDS = [
        "Ammonium dihydrogen phosphate", 
        "Extracting solvent", 
        "Diluent solution",
        "Internal standard"
    ]



    def _flush() -> None:
        nonlocal current_heading, current_lines
        if not current_heading:
            return

        sol_base = _derive_solution_name(current_heading)

        # current_lines 원소가 여러 줄을 포함할 수 있으므로 먼저 평탄화
        full_text = "\n".join(l for l in current_lines if l.strip())
        text_lines = full_text.splitlines()

        # 함량 서브섹션 분리: "30 mg/25 mg" 또는 "26 밀리그램/5 밀리그램" 줄이 있으면 각 함량별 prep 생성
        str_positions: list[tuple[int, str]] = []
        for i, l in enumerate(text_lines):
            m_en = _RE_STRENGTH_LINE.match(l)
            if m_en:
                str_positions.append((i, m_en.group(1).strip()))
                continue
            ko_str = _normalize_ko_strength(l)
            if ko_str:
                str_positions.append((i, ko_str))

        def _make_block(sol_name: str, sub_lines: list[str]) -> None:
            prep_text = "\n".join(l for l in sub_lines if l.strip())
            vol = _extract_volume(current_heading, sub_lines)
            # 내용이 없고 볼륨도 없는 빈 블록(이중언어 문서에서 영문 헤딩만 남는 경우) 제외
            if not prep_text and vol is None:
                return
            ingredients = _extract_ingredients(prep_text, vol)
            
            # 본문에 시약이 없으면 제목에서 추출 시도
            if not ingredients and current_heading:
                # 제목에서 괄호 안의 시약명 등 추출 시도 (예: Diluent (0.1mol/L Hydrochloric acid))
                heading_ingredients = _extract_ingredients(current_heading, vol)
                if heading_ingredients:
                    ingredients = heading_ingredients
                else:
                    # 제목 자체에서 시약명 유추 (예: '0.1mol/L Hydrochloric acid' 부분)
                    m = re.search(r"\(([^)]+)\)", current_heading)
                    if m:
                        ing_name = re.sub(r"^\d+(\.\d+)?\s*(mol/L|M|N)\s+", "", m.group(1), flags=re.IGNORECASE).strip()
                        # 함량 레이블("10/20/10 mg" 등)은 시약이 아님
                        if re.match(r'^[\d/.,\s]+mg\s*$', ing_name, re.IGNORECASE):
                            pass
                        # pH 조건 설명 ("pH 4.0 + 0.3 % Tween 80") 또는 약전 시험액명 (한글 포함) 제외
                        elif re.match(r'^pH\b', ing_name, re.IGNORECASE):
                            pass
                        elif re.search(r'[가-힣]', ing_name):
                            pass
                        elif len(ing_name) > 3:
                            ingredients.append({"name": ing_name, "amount": vol if vol else 0, "unit": "ml"})

            # CP029 등에서 용출 항목에 섞여 들어오는 타 항목 시약 제거

            if "dissolution" in current_heading.lower():
                ingredients = [ing for ing in ingredients if not any(kw.lower() in ing["name"].lower() for kw in EXCLUDE_KEYWORDS)]

            glassware = _extract_glassware(prep_text)
            filters_list = _extract_filters_from_text(prep_text)
            if _RE_CENTRIFUGE.search(prep_text):
                filters_list.append({
                    "size_um": None, "material": None, "manufacturer": None,
                    "filter_type": "centrifuge", "count_per_batch": 1,
                })
            blocks.append({
                "section_name": current_heading,
                "solution_name": sol_name,
                "volume_per_batch_ml": vol,
                "preparation_text": prep_text,
                "ingredients": ingredients,
                "glassware": glassware,
                "filters": filters_list,
            })

        # 첫 번째 줄이 함량 헤더일 때만 분할 (Calculation 섹션 내 함량 참조 줄 오인식 방지)
        if str_positions and str_positions[0][0] == 0:
            for k, (start_idx, str_label) in enumerate(str_positions):
                end_idx = str_positions[k + 1][0] if k + 1 < len(str_positions) else len(text_lines)
                _make_block(f"{sol_base} (for {str_label})", text_lines[start_idx + 1:end_idx])
        else:
            _make_block(sol_base, text_lines)

        current_heading = None
        current_lines = []

    for line in section_lines:
        if _is_prep_heading(line):
            _flush()
            current_heading = line.strip()
        elif _is_any_prep_heading(line):
            # 우리가 제외하기로 한 헤딩(Standard 등)이라도, 일단 헤딩이 나오면 이전 섹션은 끝내야 함
            _flush()
            current_heading = None 
        elif current_heading is not None:
            current_lines.append(line)


    _flush()
    return blocks


def _split_ko_dissolution_methods(section: dict) -> list[dict] | None:
    """
    용출 섹션 lines에 '방법 I/II/III:' 헤딩이 2개 이상이면 각 방법별 섹션으로 분리.
    없으면 None 반환.
    """
    lines = section['lines']
    boundaries: list[tuple[int, str, str]] = []
    for i, line in enumerate(lines):
        m = _RE_KO_METHOD_HEADING.match(line.strip())
        if m:
            component = (m.group(2) or '').strip()
            boundaries.append((i, m.group(1), component))

    if len(boundaries) < 2:
        return None

    result: list[dict] = []
    for k, (start_i, num, component) in enumerate(boundaries):
        end_i = boundaries[k + 1][0] if k + 1 < len(boundaries) else len(lines)
        method_lines = lines[start_i + 1:end_i]
        name = (
            f"Dissolution (방법 {num}: {component})" if component
            else f"Dissolution (방법 {num})"
        )
        result.append({'name': name, 'lines': method_lines})
    return result


# ── Body element 순서 반복 ────────────────────────────────

def _iter_body_elements(doc):
    """문서 body의 단락/표를 문서 순서대로 ('para', text) 또는 ('table', rows) yield."""
    for child in doc.element.body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            texts = [n.text or '' for n in child.iter() if n.tag == _W_T]
            yield 'para', ''.join(texts).strip()
        elif tag == 'tbl':
            rows: list[list[str]] = []
            for tr in child.iter(_W_TR):
                cells = []
                for tc in tr.iter(_W_TC):
                    ct = [n.text or '' for n in tc.iter(_W_T)]
                    cells.append(''.join(ct).strip())
                if cells:
                    rows.append(cells)
            if rows:
                yield 'table', rows


def _extract_hplc_conditions_per_section(doc) -> dict[str, dict]:
    """
    문서 body를 순서대로 스캔.
    각 시험항목 섹션에서 HPLC 조건 표(Flow rate + Run time)와
    직후 주입 순서 표(Solution | Injection No.)를 추출.
    반환: {section_name: {flow_rate_ml_min, run_time_min, injections:[...]}}
    """
    result: dict[str, dict] = {}
    current_section: str | None = None
    pending_hplc: dict | None = None

    for elem_type, elem_data in _iter_body_elements(doc):
        if elem_type == 'para':
            canonical = _get_canonical_test_item_name(elem_data)
            if canonical:
                # 섹션 전환 시 미처리 pending_hplc가 있으면 현재 섹션에 저장
                if pending_hplc is not None and current_section and current_section not in result:
                    result[current_section] = pending_hplc
                current_section = canonical
                pending_hplc = None
            elif current_section and re.match(r'^dissolution\b', current_section, re.IGNORECASE):
                # 용출 섹션 내 방법 I/II/III 헤딩 감지 → 섹션명 변경
                m_meth = _RE_KO_METHOD_HEADING.match(elem_data.strip())
                if m_meth:
                    if pending_hplc is not None and current_section not in result:
                        result[current_section] = pending_hplc
                        pending_hplc = None
                    component = (m_meth.group(2) or '').strip()
                    num = m_meth.group(1)
                    current_section = (
                        f"Dissolution (방법 {num}: {component})" if component
                        else f"Dissolution (방법 {num})"
                    )
            continue

        rows: list[list[str]] = elem_data
        if not current_section or not rows:
            continue

        row_map = {r[0]: r[1] for r in rows if len(r) >= 2}
        flow_key = next((k for k in row_map if re.match(r'^(?:flow\s*rate|유량)', k, re.IGNORECASE)), None)
        run_key  = next((k for k in row_map if re.match(r'^(?:run\s*time|분석\s*시간)', k, re.IGNORECASE)), None)

        col_key = next((k for k in row_map if re.match(r'^(?:column|칼럼)\s*$', k, re.IGNORECASE)), None)
        # 첫 행이 'Column | <spec>' 형태인 경우도 col_key로 처리
        col_spec_from_header = None
        if not col_key and len(rows) > 0 and len(rows[0]) >= 2:
            if re.match(r'^(?:column|칼럼)\s*$', rows[0][0], re.IGNORECASE):
                col_spec_from_header = rows[0][1].strip()
        if flow_key and run_key and current_section not in result:
            fm = re.search(r'(\d+(?:\.\d+)?)', row_map[flow_key])
            rm = re.search(r'(\d+(?:\.\d+)?)', row_map[run_key])
            if fm and rm:
                col_val = (row_map[col_key].strip() if col_key else col_spec_from_header)
                pending_hplc = {
                    "flow_rate_ml_min": float(fm.group(1)),
                    "run_time_min":     float(rm.group(1)),
                    "column_spec": col_val,
                    "injections": [],
                }
            continue
        # Run time 없이 Column + Flow rate만 있어도 column_spec 추출 목적으로 pending_hplc 설정
        if (col_key or col_spec_from_header) and flow_key and not run_key and current_section not in result:
            fm = re.search(r'(\d+(?:\.\d+)?)', row_map[flow_key])
            col_val = (row_map[col_key].strip() if col_key else col_spec_from_header)
            if fm and col_val:
                pending_hplc = {
                    "flow_rate_ml_min": float(fm.group(1)),
                    "run_time_min":     None,
                    "column_spec": col_val,
                    "injections": [],
                }
            continue

        if pending_hplc is not None:
            header = rows[0]
            # [Solution, Injection] 또는 [No., Solution, Injection] 형태 모두 처리
            sol_col = next((i for i, h in enumerate(header) if re.match(r'^solution', h, re.IGNORECASE)), None)
            inj_col = next((i for i, h in enumerate(header) if re.match(r'^injection', h, re.IGNORECASE)), None)
            if sol_col is not None and inj_col is not None:
                for row in rows[1:]:
                    if len(row) <= max(sol_col, inj_col):
                        continue
                    # 각주 표기 제거: "2¹⁾" → "21)" → "2"
                    raw_count = re.sub(r'\d\)\s*$', '', row[inj_col].strip()).strip()
                    count_m = re.search(r'(\d+)', raw_count) or re.search(r'(\d+)', row[inj_col])
                    count   = int(count_m.group(1)) if count_m else 1
                    pending_hplc["injections"].append({
                        "solution":          row[sol_col].strip(),
                        "count":             count,
                        "scales_with_batch": bool(re.search(r'\bsample\b|검액', row[sol_col], re.IGNORECASE)),
                    })
                result[current_section] = pending_hplc
                pending_hplc = None

    return result


# ── 용출 조건 추출 ───────────────────────────────────────

def _extract_dissolution_conditions(doc) -> dict | None:
    """
    문서 body 표에서 용출 조건(Volume, 배지, 장치, 속도) 파싱.
    Dissolution Media / Volume 두 행이 모두 있는 표를 탐색한다.
    """
    for tbl in doc.tables:
        row_map: dict[str, str] = {}
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2:
                row_map[cells[0].strip()] = cells[1].strip()

        has_diss_media = any(
            re.match(r"dissolution\s+medi", k, re.IGNORECASE) for k in row_map
        )
        # "Volume" 또는 "Vessel Volume" 모두 허용
        has_volume = any(re.search(r"\bvolume\b", k, re.IGNORECASE) for k in row_map)
        # 한글 용출 조건 표: "용량" + "장치" 또는 "속도"
        has_ko_volume = any(re.match(r'^용량$', k) for k in row_map)
        has_ko_apparatus = any(re.match(r'^(?:장치|속도)', k) for k in row_map)
        is_ko_diss = has_ko_volume and has_ko_apparatus

        if not ((has_diss_media and has_volume) or is_ko_diss):
            continue

        cond: dict = {"vessels_per_batch": 6}
        for raw_key, val in row_map.items():
            # 영문 키
            if re.match(r"dissolution\s+medi", raw_key, re.IGNORECASE):
                cond["medium_name"] = val
            elif re.search(r"\bvolume\b", raw_key, re.IGNORECASE) and "volume_per_vessel_ml" not in cond:
                m = re.search(r"(\d+(?:\.\d+)?)\s*mL", val, re.IGNORECASE)
                if m:
                    cond["volume_per_vessel_ml"] = float(m.group(1))
            elif re.match(r"(?:usp\s+)?apparatus\b", raw_key, re.IGNORECASE):
                cond["apparatus"] = val
            elif re.match(r"speed\b", raw_key, re.IGNORECASE):
                ms = re.search(r"(\d+)", val)
                if ms:
                    cond["speed_rpm"] = int(ms.group(1))
            elif re.match(r"sampling\s+time", raw_key, re.IGNORECASE):
                cond["sampling_time"] = val
            # 한글 키
            elif re.match(r'^용량$', raw_key):
                m = re.search(r"(\d+(?:\.\d+)?)\s*mL", val, re.IGNORECASE)
                if m:
                    cond["volume_per_vessel_ml"] = float(m.group(1))
            elif re.match(r'^시험액$', raw_key):
                cond["medium_name"] = val
            elif re.match(r'^장치', raw_key):
                cond["apparatus"] = val
            elif re.match(r'^속도', raw_key):
                ms = re.search(r"(\d+)", val)
                if ms:
                    cond["speed_rpm"] = int(ms.group(1))
            elif re.match(r'^샘플링', raw_key):
                cond["sampling_time"] = val

        if is_ko_diss and not cond.get("medium_name"):
            cond["medium_name"] = "시험액"

        if "volume_per_vessel_ml" in cond:
            return cond

    return None


def _extract_dissolution_conditions_per_method(doc) -> dict[str, dict]:
    """
    문서 body를 순서대로 스캔해 방법별(방법 I/II/III) 용출 조건 추출.
    반환: {section_name: dissolution_conditions_dict}
    """
    result: dict[str, dict] = {}
    current_section: str | None = None

    def _parse_cond_table(rows: list[list[str]]) -> dict | None:
        row_map = {r[0].strip(): r[1].strip() for r in rows if len(r) >= 2}
        has_diss_media = any(re.match(r"dissolution\s+medi", k, re.IGNORECASE) for k in row_map)
        has_volume     = any(re.search(r"\bvolume\b", k, re.IGNORECASE) for k in row_map)
        has_ko_volume  = any(re.match(r'^용량$', k) for k in row_map)
        has_ko_app     = any(re.match(r'^(?:장치|속도)', k) for k in row_map)
        is_ko_diss     = has_ko_volume and has_ko_app
        if not ((has_diss_media and has_volume) or is_ko_diss):
            return None
        cond: dict = {"vessels_per_batch": 6}
        for raw_key, val in row_map.items():
            if re.match(r"dissolution\s+medi", raw_key, re.IGNORECASE):
                cond["medium_name"] = val
            elif re.search(r"\bvolume\b", raw_key, re.IGNORECASE) and "volume_per_vessel_ml" not in cond:
                mv = re.search(r"(\d+(?:\.\d+)?)\s*mL", val, re.IGNORECASE)
                if mv:
                    cond["volume_per_vessel_ml"] = float(mv.group(1))
            elif re.match(r"(?:usp\s+)?apparatus\b", raw_key, re.IGNORECASE):
                cond["apparatus"] = val
            elif re.match(r"speed\b", raw_key, re.IGNORECASE):
                ms = re.search(r"(\d+)", val)
                if ms:
                    cond["speed_rpm"] = int(ms.group(1))
            elif re.match(r"sampling\s+time", raw_key, re.IGNORECASE):
                cond["sampling_time"] = val
            elif re.match(r'^용량$', raw_key):
                mv = re.search(r"(\d+(?:\.\d+)?)\s*mL", val, re.IGNORECASE)
                if mv:
                    cond["volume_per_vessel_ml"] = float(mv.group(1))
            elif re.match(r'^시험액$', raw_key):
                cond["medium_name"] = val
            elif re.match(r'^장치', raw_key):
                cond["apparatus"] = val
            elif re.match(r'^속도', raw_key):
                ms = re.search(r"(\d+)", val)
                if ms:
                    cond["speed_rpm"] = int(ms.group(1))
            elif re.match(r'^샘플링', raw_key):
                cond["sampling_time"] = val
        if is_ko_diss and not cond.get("medium_name"):
            cond["medium_name"] = "시험액"
        return cond if "volume_per_vessel_ml" in cond else None

    for elem_type, elem_data in _iter_body_elements(doc):
        if elem_type == 'para':
            canonical = _get_canonical_test_item_name(elem_data)
            if canonical:
                current_section = canonical
            elif current_section and re.match(r'^dissolution\b', current_section, re.IGNORECASE):
                m_meth = _RE_KO_METHOD_HEADING.match(elem_data.strip())
                if m_meth:
                    component = (m_meth.group(2) or '').strip()
                    num = m_meth.group(1)
                    current_section = (
                        f"Dissolution (방법 {num}: {component})" if component
                        else f"Dissolution (방법 {num})"
                    )
            continue
        # 표 처리
        if not current_section:
            continue
        if not re.match(r'^dissolution\b', current_section, re.IGNORECASE):
            continue
        if current_section in result:
            continue
        cond = _parse_cond_table(elem_data)
        if cond:
            result[current_section] = cond

    return result


def _extract_standard_medium_ml_for_dissolution(
    dissolution_section_lines: list[str],
    product_strengths: list[str],
) -> dict[str, float]:
    """
    용출 시험 섹션에서 함량별 표준액 조제 시 필요한 시험액(dissolution medium) 총량 추출.

    규칙:
    - 'standard'를 포함하지만 'sample'은 포함하지 않는 preparation 블록만 대상
    - dissolution medium을 희석제로 쓰는 블록 안의 volumetric flask 크기를 합산
    - 함량별 레이블(예: "30 mg/25 mg")이 블록 내에 있으면 함량별로 분리
    - 헤딩에 "(for X mg)" 패턴이 있으면 해당 함량에 귀속
    반환: {strength_label: total_ml}
    """
    # 준비 블록 분리
    blocks: list[tuple[str, list[str]]] = []
    cur_heading: str | None = None
    cur_lines: list[str] = []
    for line in dissolution_section_lines:
        if _is_prep_heading(line):
            if cur_heading is not None:
                blocks.append((cur_heading, cur_lines[:]))
            cur_heading = line.strip()
            cur_lines = []
        elif cur_heading is not None:
            cur_lines.append(line)
    if cur_heading is not None:
        blocks.append((cur_heading, cur_lines))

    per_strength: dict[str, float] = {}
    no_strength_total: float = 0.0

    for heading, lines in blocks:
        is_standard = (
            re.search(r"\bstandard\b", heading, re.IGNORECASE) or
            re.search(r"표준", heading)
        )
        is_sample = (
            re.search(r"\bsample\b", heading, re.IGNORECASE) or
            re.search(r"검액", heading)
        )
        if not is_standard or is_sample:
            continue

        # 헤딩의 "(for X mg)" 패턴
        heading_strength: str | None = None
        m_hs = _RE_STRENGTH_PAREN.search(heading)
        if m_hs:
            heading_strength = m_hs.group(1).strip()

        # 블록 내 함량 레이블로 서브 블록 분리 (영문 "30 mg" + 한글 "26 밀리그램/5 밀리그램")
        sub_blocks: list[tuple[str | None, list[str]]] = []
        cur_sub_str: str | None = None
        cur_sub_lines: list[str] = []
        found_str_in_lines = False

        for line in lines:
            m_sl = _RE_STRENGTH_LINE.match(line)
            m_ko = _RE_KO_STRENGTH_LINE.match(line) if not m_sl else None
            if m_sl:
                if cur_sub_lines:
                    sub_blocks.append((cur_sub_str, cur_sub_lines[:]))
                cur_sub_str = m_sl.group(1).strip()
                cur_sub_lines = []
                found_str_in_lines = True
            elif m_ko:
                if cur_sub_lines:
                    sub_blocks.append((cur_sub_str, cur_sub_lines[:]))
                cur_sub_str = _normalize_ko_strength(line)
                cur_sub_lines = []
                found_str_in_lines = True
            else:
                cur_sub_lines.append(line)
        if cur_sub_lines:
            sub_blocks.append((cur_sub_str, cur_sub_lines))

        if not found_str_in_lines:
            sub_blocks = [(heading_strength, lines)]

        for sub_strength, sub_lines in sub_blocks:
            combined = heading + " " + " ".join(sub_lines)
            # 영문/한글 dissolution medium BTV 감지
            if not re.search(
                r"\bdissolution\s+(?:medium|media)\b|시험액",
                combined, re.IGNORECASE
            ):
                continue

            # volumetric flask 크기 합산 (영문 + 한글)
            vol_sum = sum(float(m.group(1)) for m in _RE_VOLUMETRIC_FLASK_SIZE.finditer(combined))
            vol_sum += sum(float(m.group(1)) for m in _RE_KO_VOLUMETRIC_FLASK_SIZE.finditer(combined))
            if vol_sum <= 0:
                vol_sum = _extract_volume(heading, sub_lines) or 0.0
            if vol_sum <= 0:
                continue

            if sub_strength:
                per_strength[sub_strength] = per_strength.get(sub_strength, 0.0) + vol_sum
            else:
                no_strength_total += vol_sum

    if per_strength:
        return per_strength
    if no_strength_total > 0:
        return {s: no_strength_total for s in product_strengths}
    return {}


# ── Standard 표 파싱 ─────────────────────────────────────

def _extract_standards_per_section(doc) -> dict[str, list[dict]]:
    """
    문서 body를 순서대로 스캔해 각 시험항목 섹션의 STD Name | Grade 표를 추출.
    반환: {section_name: [{"std_name": ..., "grade": ...}]}
    """
    result: dict[str, list[dict]] = {}
    current_section: str | None = None

    for elem_type, elem_data in _iter_body_elements(doc):
        if elem_type == 'para':
            canonical = _get_canonical_test_item_name(elem_data)
            if canonical:
                parent = _parent_section_name(canonical, current_section)
                if not parent:
                    current_section = canonical
            elif current_section and re.match(r'^dissolution\b', current_section, re.IGNORECASE):
                # 용출 섹션 내 방법 I/II/III 헤딩 감지
                m_meth = _RE_KO_METHOD_HEADING.match(elem_data.strip())
                if m_meth:
                    component = (m_meth.group(2) or '').strip()
                    num = m_meth.group(1)
                    current_section = (
                        f"Dissolution (방법 {num}: {component})" if component
                        else f"Dissolution (방법 {num})"
                    )
            continue

        # 표 처리
        rows: list[list[str]] = elem_data
        if not current_section or not rows:
            continue

        header = rows[0]
        std_name_col = next(
            (i for i, h in enumerate(header) if re.match(r'(?:std(?:\s*name)?|standard(?:\s+name)?|표준품)', h, re.IGNORECASE)),
            None,
        )
        if std_name_col is None:
            continue

        grade_col = next(
            (i for i, h in enumerate(header) if re.match(r'(?:grade|등급)', h, re.IGNORECASE)),
            None,
        )

        standards: list[dict] = []
        for row in rows[1:]:
            std_name = row[std_name_col].strip() if std_name_col < len(row) else ""
            grade = (
                row[grade_col].strip()
                if grade_col is not None and grade_col < len(row)
                else ""
            )
            if std_name:
                standards.append({"std_name": std_name, "grade": grade})

        if standards and current_section not in result:
            result[current_section] = standards

    return result


# ── Reagent 표 파싱 ──────────────────────────────────────

def _extract_all_reagents(doc) -> list[dict]:
    """
    문서의 모든 Reagent 표(헤더 행: Reagent | Grade | Manufacturer | Cat. No. | Tracking No.)
    에서 시약 정보를 추출해 반환.
    """
    reagents: list[dict] = []
    seen: set[tuple] = set()

    for tbl in doc.tables:
        rows = [[c.text.strip() for c in row.cells] for row in tbl.rows]
        if not rows or len(rows[0]) < 2:
            continue

        # 첫 행이 Reagent 또는 한글 시약 헤더인지 확인
        header = rows[0]
        if not re.match(r"^(?:reagent|시약)\s*$", header[0], re.IGNORECASE):
            continue

        # 컬럼 인덱스 매핑 (영문/한글 공통)
        col: dict[str, int] = {}
        for i, h in enumerate(header):
            hl = h.lower()
            if re.match(r"^(?:reagent|시약)$", hl):
                col["name"] = i
            elif "tracking" in hl or "추적" in hl:
                col["tracking_no"] = i
            elif "grade" in hl or "등급" in hl:
                col["grade"] = i
            elif "manufacturer" in hl or "제조처" in hl:
                col["manufacturer"] = i
            elif "cat" in hl:
                col["cat_no"] = i

        if "name" not in col or "tracking_no" not in col:
            continue

        for row in rows[1:]:
            name = row[col["name"]].strip() if col["name"] < len(row) else ""
            tracking = row[col["tracking_no"]].strip() if col["tracking_no"] < len(row) else ""
            if not name or not tracking:
                continue
            key = (name.lower(), tracking)
            if key in seen:
                continue
            seen.add(key)
            reagents.append({
                "name": name,
                "grade": row[col["grade"]].strip() if col.get("grade") is not None and col["grade"] < len(row) else "",
                "manufacturer": row[col["manufacturer"]].strip() if col.get("manufacturer") is not None and col["manufacturer"] < len(row) else "",
                "cat_no": row[col["cat_no"]].strip() if col.get("cat_no") is not None and col["cat_no"] < len(row) else "",
                "tracking_no": tracking,
            })

    return reagents


def _build_reagent_lookup(reagents: list[dict]) -> dict[str, list[str]]:
    """시약 이름(소문자) → tracking_no 리스트 (중복 제거)."""
    lookup: dict[str, list[str]] = {}
    for r in reagents:
        key = r["name"].lower()
        if key not in lookup:
            lookup[key] = []
        if r["tracking_no"] not in lookup[key]:
            lookup[key].append(r["tracking_no"])
    return lookup


def _enrich_ingredients_tracking(
    test_items: list[dict],
    reagent_lookup: dict[str, list[str]],
) -> None:
    """ingredients 각 항목에 tracking_no 필드를 추가 (일치하는 경우)."""
    for item in test_items:
        for prep in item.get("preparations", []):
            for ing in prep.get("ingredients", []):
                key = ing["name"].lower()
                tracking = reagent_lookup.get(key)
                if tracking:
                    ing["tracking_no"] = ", ".join(tracking)


# ── 헤더 표에서 제품명·함량 추출 ─────────────────────────

def _extract_header_title(doc) -> tuple[str | None, str | None]:
    """
    Word 섹션 헤더 표의 'Title: ...' 셀에서 (제품명, 함량)을 추출.
    예) 'Title: CTPH-D005 5 mg (Vonoprazan fumarate)'
         → ('CTPH-D005', '5 mg')
    예) 'Title: Pioglitazone Hydrochloride/Empagliflozin tablets'
         → ('Pioglitazone Hydrochloride/Empagliflozin tablets', None)
    """
    for section in doc.sections:
        for tbl in section.header.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    m = re.match(r"Title:\s*(.+)", text, re.IGNORECASE)
                    if not m:
                        continue
                    # 한글 줄 제거: 첫 번째 줄(영문)만 사용
                    first_line = m.group(1).split("\n")[0].strip()
                    # 함량 추출: "5 mg", "30 mg/25 mg", "25/10mg" 형식 모두 지원
                    strength_m = re.search(
                        r"\b(\d+(?:\.\d+)?\s*mg(?:\s*/\s*\d+(?:\.\d+)?\s*mg)*"
                        r"|\d+(?:\.\d+)?/\d+(?:\.\d+)?\s*mg)\b",
                        first_line, re.IGNORECASE
                    )
                    strength = strength_m.group(1).strip() if strength_m else None
                    if strength:
                        # "10mg" → "10 mg", "25/10mg" → "25/10 mg"
                        strength = re.sub(r'(\d)(mg)', r'\1 \2', strength, flags=re.IGNORECASE)
                    # 제품명: 괄호 앞 부분, 함량 제거
                    name = re.sub(r"\s*\([^)]*\)\s*$", "", first_line).strip()
                    if strength:
                        # 함량 숫자를 제품명에서 제거 ("CTPH-D005 5 mg" → "CTPH-D005")
                        name = re.sub(
                            r"\s*\d+(?:\.\d+)?(?:/\d+(?:\.\d+)?)?\s*mg(?:\s*/\s*\d+(?:\.\d+)?\s*mg)*\s*",
                            " ", name, flags=re.IGNORECASE
                        ).strip()
                    return name or None, strength
    return None, None


def _extract_doc_no(doc) -> str | None:
    """헤더 표에서 STM 문서 번호 추출 (예: STM-300599-T4)."""
    for section in doc.sections:
        for tbl in section.header.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    m = re.match(r"^(STM-\S+)", text, re.IGNORECASE)
                    if m:
                        return m.group(1)
    return None


# ── 문서 파싱 진입점 ──────────────────────────────────────

def parse_document(doc_path: str) -> dict:
    """단일 STM .docx 파일을 파싱해 구조화된 dict 반환."""
    doc = Document(doc_path)

    # 헤더 표에서 제품명·함량 우선 추출
    header_name, header_strength = _extract_header_title(doc)

    # 전체 단락 추출 (한글 문서 대응)
    all_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    english_lines = [l for l in all_lines if not _is_korean(l)]

    # 제품명: 헤더 우선, 없으면 본문 파싱
    if header_name:
        product_name = header_name
    else:
        names = _extract_product_names(english_lines)
        product_name = " / ".join(names) if names else Path(doc_path).stem

    # 함량: 헤더에서 찾았으면 사용, 없으면 본문에서 추출 (한글 포함)
    if header_strength:
        strengths = [header_strength]
    else:
        strengths = _extract_strengths(all_lines)

    # 국영문(bilingual) 감지: 한글·영문 줄이 모두 10개 이상이면 영문만 파싱
    _ko_cnt = sum(1 for l in all_lines if _is_korean(l))
    _en_cnt = len(all_lines) - _ko_cnt
    _is_bilingual = _ko_cnt > 10 and _en_cnt > 10
    parse_lines = english_lines if _is_bilingual else all_lines

    # 시험항목 섹션 파싱
    sections = _parse_sections(parse_lines)

    # 한글 용출 방법 분리 (방법 I/II/III)
    expanded_sections: list[dict] = []
    for sec in sections:
        if re.match(r'^dissolution\b', sec['name'], re.IGNORECASE):
            sub = _split_ko_dissolution_methods(sec)
            if sub:
                expanded_sections.extend(sub)
            else:
                expanded_sections.append(sec)
        else:
            expanded_sections.append(sec)
    sections = expanded_sections

    test_items: list[dict] = []
    for sec in sections:
        preps = _parse_preparations(sec["lines"])
        test_items.append({"name": sec["name"], "preparations": preps})

    # 용출 시험 조건 추출 (방법별)
    diss_conds_by_section = _extract_dissolution_conditions_per_method(doc)
    # 단일 dissolution 섹션 fallback (기존 호환)
    if not diss_conds_by_section:
        single = _extract_dissolution_conditions(doc)
        if single:
            diss_conds_by_section = {"Dissolution": single}

    for item in test_items:
        if not re.match(r"^dissolution\b", item["name"], re.IGNORECASE):
            continue
        cond = diss_conds_by_section.get(item["name"])
        if cond is None:
            # fallback: 첫 번째 조건 사용
            cond = next(iter(diss_conds_by_section.values()), None)
        if cond:
            sec = next((s for s in sections if s["name"] == item["name"]), None)
            method_conds = dict(cond)
            method_conds["standard_medium_ml_by_strength"] = (
                _extract_standard_medium_ml_for_dissolution(
                    sec["lines"] if sec else [], strengths
                )
            )
            item["dissolution_conditions"] = method_conds

    # Reagent 표 파싱 → ingredients에 tracking_no 추가
    reagents = _extract_all_reagents(doc)
    reagent_lookup = _build_reagent_lookup(reagents)
    _enrich_ingredients_tracking(test_items, reagent_lookup)

    # 표준품(STD Name) 표 추출 → test_items에 연결
    standards_per_section = _extract_standards_per_section(doc)
    for item in test_items:
        if item["name"] in standards_per_section:
            item["standards"] = standards_per_section[item["name"]]

    # HPLC 크로마토그래피 조건 추출 (Mobile phase 볼륨 계산용)
    hplc_per_section = _extract_hplc_conditions_per_section(doc)
    for item in test_items:
        if item["name"] in hplc_per_section:
            item["hplc_conditions"] = hplc_per_section[item["name"]]

    # 문서 번호
    doc_no = _extract_doc_no(doc)

    # ── CP001 전용 보정 로직 (사용자 요청 반영) ───────────────────
    if "CP001" in Path(doc_path).name.upper():
        product_name = "NesinaMet Tablets"
        strengths = ["12.5/500 mg"]
        
        for item in test_items:
            # 1. 확인시험 조제 정보 수정
            if "Identification" in item["name"]:
                for p in item["preparations"]:
                    if "Diluent" in p["solution_name"]:
                        p["volume_per_batch_ml"] = 75.0
                        p["preparation_text"] = "염산 8.5 mL을 물에 넣어 1000 mL로 한다."
                        p["ingredients"] = [
                            {"name": "Hydrochloric acid", "amount": 0.6375, "unit": "ml"}
                        ]
                        p["glassware"] = [] 
                        p["is_reagent"] = True

            # 2. 용출시험 조제 정보 수정 (필수 용액 유지: 시험액, 이동상, 표준액, 검액)
            if "Dissolution" in item["name"]:
                # CP001 용출시험에 필요한 용액 리스트
                allowed_names = [
                    "시험액 (Dissolution medium)",
                    "Buffer",
                    "Mobile phase",
                    "Alogliptin standard stock solution",
                    "Metformin Hydrochloride standard stock solution",
                    "Standard solution",
                    "Sample solution"
                ]
                filtered_preps = []
                seen_names = set()
                
                for p in item["preparations"]:
                    sol_name = p["solution_name"]
                    
                    # (1) 시험액 보정
                    if any(kw in sol_name for kw in ["Dissolution medium", "시험액"]):
                        p["solution_name"] = "시험액 (Dissolution medium)"
                        p["volume_per_batch_ml"] = 1000.0
                        p["preparation_text"] = "Ammonium phosphate monobasic 2.3 g을 물 1000 mL에 녹이고 인산으로 pH 6.0으로 맞춘다."
                        p["ingredients"] = [
                            {"name": "Ammonium phosphate monobasic", "amount": 2.3, "unit": "g"}
                        ]
                        p["is_reagent"] = True
                        p["glassware"] = []
                    
                    # (2) 이동상 조제용 완충액 보정 (이름 변경: Buffer)
                    if "Ammonium dihydrogen phosphate" in sol_name:
                        p["solution_name"] = "Buffer"
                        p["volume_per_batch_ml"] = 1000.0
                        p["preparation_text"] = "Ammonium dihydrogen phosphate 5.75 g을 물에 녹여 1000 mL로 한다."
                        p["ingredients"] = [
                            {"name": "Ammonium dihydrogen phosphate", "amount": 5.75, "unit": "g"}
                        ]
                        p["is_reagent"] = True
                        p["glassware"] = []

                    # (3) 이동상 보정 (Buffer 이름 반영)
                    if "Mobile phase" in sol_name:
                        p["volume_per_batch_ml"] = 1000.0
                        p["preparation_text"] = "Mix Buffer and acetonitrile in the ratio of 500 : 500 v/v. Then add 7.2 g of Sodium dodecyl sulfate."
                        p["ingredients"] = [
                            {"name": "Buffer", "amount": 500.0, "unit": "ml"},
                            {"name": "Acetonitrile", "amount": 500.0, "unit": "ml"},
                            {"name": "Sodium dodecyl sulfate", "amount": 7.2, "unit": "g"}
                        ]
                        p["is_reagent"] = True
                        p["glassware"] = []

                    # (4) 검액 필터 정보 추가 및 텍스트 보정
                    if "Sample solution" in sol_name:
                        p["preparation_text"] = "위의 용출 조건에 따라 시험을 실시하고, 각 시간별로 검액 10 mL를 취하여 0.45 um PVDF 또는 Nylon 필터로 여과한다."
                        p["filters"] = [
                            {"size_um": 0.45, "material": "PVDF", "manufacturer": "Hyundai micro", "filter_type": "syringe", "count_per_batch": 1},
                            {"size_um": 0.45, "material": "Nylon", "manufacturer": "Hyundai micro", "filter_type": "syringe", "count_per_batch": 1}
                        ]
                    
                    # 허용된 리스트 필터링 및 중복 제거
                    if (p["solution_name"] in allowed_names or "Ammonium dihydrogen phosphate" in p["solution_name"]) and p["solution_name"] not in seen_names:
                        if "extracting" in (p.get("preparation_text") or "").lower():
                            continue
                        filtered_preps.append(p)
                        seen_names.add(p["solution_name"])
                
                item["preparations"] = filtered_preps
                
                # 용출 조건 테이블 정보 보정
                if "dissolution_conditions" not in item:
                    item["dissolution_conditions"] = {
                        "vessels_per_batch": 6,
                        "apparatus": "USP-II (Paddle)",
                        "speed_rpm": 50,
                        "sampling_time": "30 minutes"
                    }
                item["dissolution_conditions"]["medium_name"] = "시험액 (pH 6.0)"
                item["dissolution_conditions"]["volume_per_vessel_ml"] = 900.0
                item["dissolution_conditions"]["standard_medium_ml_by_strength"] = {
                    "12.5/500 mg": 300.0
                }

 




    return {
        "id": Path(doc_path).stem,
        "doc_no": doc_no,
        "stm_file": Path(doc_path).name,
        "product_name": product_name,
        "strengths": strengths,
        "test_items": test_items,
    }


def build_knowledge_base(progress_callback=None) -> list[dict]:
    """STM 폴더의 모든 .docx 파일을 파싱해 knowledge_base.json 저장."""
    stm_folder = _NETWORK_STM if _NETWORK_STM.exists() else _LOCAL_STM
    products: list[dict] = []
    files = sorted(f for f in stm_folder.glob("*.docx") if not f.name.startswith("~$"))

    for i, doc_path in enumerate(files):
        msg = f"Parsing {doc_path.name} ({i+1}/{len(files)})..."
        print(f"  → {msg}")
        if progress_callback:
            progress_callback(msg)
        product = parse_document(str(doc_path))
        products.append(product)

    DATA_FOLDER.mkdir(exist_ok=True)
    KB_PATH.write_text(
        json.dumps({"products": products}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(f"Knowledge base saved → {KB_PATH}")
    return products


def load_knowledge_base() -> list[dict]:
    if KB_PATH.exists():
        data = json.loads(KB_PATH.read_text(encoding="utf-8"))
        return data.get("products", [])
    return []


def save_knowledge_base(products: list[dict]) -> None:
    DATA_FOLDER.mkdir(exist_ok=True)
    KB_PATH.write_text(
        json.dumps({"products": products}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
