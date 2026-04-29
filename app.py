import logging
import streamlit as st
import pandas as pd
from docx import Document
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
import re
import io
import numpy as np
import os

_log = logging.getLogger(__name__)

# ── 한글 폰트 설정 (Windows Malgun Gothic) ───────────────────────────────────
try:
    from matplotlib import font_manager as fm
    font_path = r'C:\Windows\Fonts\malgun.ttf'
    if os.path.exists(font_path):
        fm.fontManager.addfont(font_path)
        plt.rcParams['font.family'] = 'Malgun Gothic'
    else:
        plt.rcParams['font.family'] = ['Malgun Gothic', 'DejaVu Sans']
except Exception:
    plt.rcParams['font.family'] = ['Malgun Gothic', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

LINE_COLORS = ['#1f4e79', '#c55a11', '#375623', '#843c0c',
               '#7030a0', '#4472c4', '#ed7d31', '#a9d18e']


# ── 유틸 함수 ─────────────────────────────────────────────────────────────────

def make_unique(cols):
    counts = {}
    result = []
    for col in cols:
        s = str(col)
        if s in counts:
            counts[s] += 1
            result.append(f"{s}.{counts[s]}")
        else:
            counts[s] = 0
            result.append(s)
    return result


def parse_time_cols(columns):
    """Initial / 숫자M / Month 패턴의 열 이름 반환."""
    return [c for c in columns
            if re.search(r'INITIAL|\b\d+M\b|MONTH', str(c).upper())]


def parse_value(val):
    """그래프용 숫자 변환. N.D / < D.L → 0.0, 없으면 None."""
    if pd.isna(val):
        return None
    s = str(val).strip().upper()
    if s in {'< D.L', '<D.L', '< D.L.', 'N.D', 'ND', 'N.D.', 'NOT DETECTED',
             '< DL', 'DISREGARD LIMIT', '≤ DISREGARD LIMIT'}:
        return 0.0
    nums = re.findall(r'-?\d+\.?\d*', s)
    return float(nums[0]) if nums else None


def parse_limit_from_text(text):
    """(limit_min, limit_max) 파싱. 범위 또는 상한만 인식."""
    # 94.0 %~105.0 %
    m = re.search(r'(\d+\.?\d*)\s*%?\s*[~～]\s*(\d+\.?\d*)\s*%', text)
    if m:
        return float(m.group(1)), float(m.group(2))
    # ≤ 1.7 %
    m = re.search(r'[≤<]\s*=?\s*(\d+\.?\d*)\s*%', text)
    if m:
        return None, float(m.group(1))
    return None, None


def extract_batch_from_context(paras):
    """주변 단락에서 배치번호(예: CATX001KRA) 추출 시도."""
    for para in reversed(paras):
        m = re.search(r'\b([A-Z]{2,6}\d{3,}[A-Z]{0,6})\b', para)
        if m:
            return m.group(1)
        m = re.search(
            r'(?:배치|Batch|LOT|Lot)\s*(?:번호|No\.?)?\s*[:\-]?\s*([A-Za-z0-9\-]+)',
            para, re.IGNORECASE)
        if m:
            return m.group(1)
    return ""


# ── 문서 파싱 ─────────────────────────────────────────────────────────────────

def _clean(cell_text: str) -> str:
    """셀 텍스트의 모든 공백·줄바꿈 정규화."""
    return ' '.join(cell_text.split())


def extract_batch_from_source1(doc):
    """
    [Source 1]의 3행 4열(Batch No / Batch Size) 데이터 참조.
    '/' 기호를 구분자로 하여 왼쪽 텍스트 추출, 없으면 전체 텍스트 (공백 제거).
    """
    try:
        if len(doc.tables) > 0:
            tbl = doc.tables[0]
            if len(tbl.rows) > 2:
                row = tbl.rows[2]
                if len(row.cells) > 3:
                    raw_text = row.cells[3].text
                    text_no_space = re.sub(r'\s+', '', raw_text)
                    if not text_no_space:
                        return None
                    if '/' in text_no_space:
                        return text_no_space.split('/')[0]
                    return text_no_space
    except Exception:
        pass
    return None


def process_document(file):
    """
    확정된 표 구조 기반 파싱.

    표 위치  : doc.tables[1] (두 번째 표, index 1)
    헤더 행  : 두 번째 행 (raw_data[1]) — Initial, 3M 등 시점 헤더 위치
    데이터 행: raw_data[2] 이후

    열 인덱스 (절대 고정)
      index 0 : No./순번       → 제외
      index 1 : 대분류          → '함량'/'유연물질' 보조 확인 + Forward Fill
      index 2 : 세부항목/소분류 → 키워드 메인 확인 + Forward Fill
      index 3 : 허용기준(spec)  → 그대로 limit_text 수집
      index 4+ : 시점 데이터   → Initial, 3M, 6M … 전부 수집

    Returns
    -------
    assay_rows : list of dict
    imp_groups : dict {차트키: [dict, ...]}
    imp_limits : dict {차트키: limit_text}
    diag       : dict  진단용 raw 정보
    global_batch_name : str 추출된 배치 번호
    """
    doc = Document(file)
    global_batch_name = extract_batch_from_source1(doc)

    assay_rows = []
    imp_groups = {}
    imp_limits = {}
    imp_row_counter: dict = {}
    diag: dict = {'n_tables': len(doc.tables), 'raw_header': [], 'raw_rows': []}

    # ── 대상 표 선택 ──────────────────────────────────────────────────────────
    if len(doc.tables) >= 2:
        target_tables = doc.tables[1:]   # index 1부터 끝까지 (복수 배치 대응)
    elif len(doc.tables) == 1:
        target_tables = doc.tables       # fallback
    else:
        return assay_rows, imp_groups, imp_limits, diag, global_batch_name

    batch_counter = [0]

    for tbl in target_tables:
        # 모든 셀: 공백·줄바꿈 완전 제거
        raw_data = [
            [_clean(cell.text) for cell in row.cells]
            for row in tbl.rows
        ]

        if len(raw_data) < 3:   # 타이틀행 + 헤더행 + 데이터행 최소 3행 필요
            continue

        # ── 헤더: 두 번째 행(index 1) ────────────────────────────────────────
        headers = raw_data[1]
        data_rows = raw_data[2:]
        diag['raw_header'] = headers
        diag['raw_rows'] = data_rows[:5]   # 진단용: 처음 5행만

        if len(headers) <= 4:
            continue

        # 시점 열 이름: index 4부터 (빈 헤더는 제외)
        time_col_names = [h for h in headers[4:] if h]
        if not time_col_names:
            continue
        n_time = len(time_col_names)

        # 배치명 추출 (우선순위: 명시된 Source 1 → 표 타이틀행 → 단락 → 표 전체 스캔)
        all_paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        batch_name = global_batch_name or (
            extract_batch_from_context([' '.join(raw_data[0])])       # 표 1행
            or extract_batch_from_context([' '.join(headers)])         # 표 2행(헤더)
            or extract_batch_from_context(all_paras[:40])              # 단락
            or extract_batch_from_context(                             # 표 전체
                [' '.join(cell for row in raw_data for cell in row)]
            )
        )
        if not batch_name:
            batch_counter[0] += 1
            batch_name = f"Batch {batch_counter[0]}"

        last_col1 = ""   # index 1 forward fill
        last_col2 = ""   # index 2 forward fill

        for row_data in data_rows:
            if len(row_data) <= 4:
                continue

            col1 = row_data[1]   # 대분류 (이미 _clean 적용됨)
            col2 = row_data[2]   # 세부항목

            # ── Forward Fill ─────────────────────────────────────────────────
            if col1 and col1.lower() != 'nan':
                last_col1 = col1
            if col2 and col2.lower() != 'nan':
                last_col2 = col2

            # 분류 키워드: index 2(세부) 우선, 없으면 index 1(대분류) 참조
            current_item = last_col2 or last_col1
            if not current_item:
                continue

            # 허용기준: index 3 고정
            limit_text = row_data[3] if len(row_data) > 3 else ""
            if limit_text.lower() == 'nan':
                limit_text = ""

            # 시점 데이터: index 4부터 n_time개
            values = row_data[4: 4 + n_time]
            values += [''] * (n_time - len(values))   # 부족분 패딩

            row_series = pd.Series(dict(zip(time_col_names, values)))

            item_upper = current_item.upper()
            entry = {
                'batch':      batch_name,
                'row':        row_series,
                'time_cols':  time_col_names,
                'limit_text': limit_text,
                'item_name':  current_item,
            }

            # ── 분류 ─────────────────────────────────────────────────────────
            if '함량' in current_item or 'ASSAY' in item_upper:
                assay_rows.append(entry)

            elif ('유연물질' in current_item or '불순물' in current_item
                  or 'IMPURIT' in item_upper or 'RELATED SUBSTANCE' in item_upper):

                # 차트 키 우선순위:
                # ① index 2가 구체적 이름 → 그대로
                # ② index 1이 구체적 이름 → 그대로
                # ③ 허용기준으로 구분
                # ④ 순번 부여
                generic = {'유연물질', 'related substances', 'impurities', 'impurity'}
                if last_col2 and last_col2.lower() not in generic:
                    imp_key = last_col2
                elif last_col1 and last_col1.lower() not in generic:
                    imp_key = last_col1
                elif limit_text:
                    imp_key = f"{current_item} ({limit_text})"
                else:
                    cnt = imp_row_counter.get(current_item, 0) + 1
                    imp_row_counter[current_item] = cnt
                    imp_key = f"{current_item} {cnt}"

                imp_groups.setdefault(imp_key, []).append(entry)
                if imp_key not in imp_limits and limit_text:
                    imp_limits[imp_key] = limit_text

    return assay_rows, imp_groups, imp_limits, diag, global_batch_name


def build_dataframe(entries):
    """
    entry 목록 → (DataFrame, time_cols).
    DataFrame 구조: 1열=배치번호, 나머지=시점 열.
    """
    if not entries:
        return None, []

    # 시점 열 순서 통합 (첫 등장 순 유지)
    seen = set()
    all_time = []
    for e in entries:
        for c in e['time_cols']:
            if c not in seen:
                all_time.append(c)
                seen.add(c)

    records = []
    for e in entries:
        rec = {'배치번호': e['batch']}
        for c in all_time:
            rec[c] = e['row'][c] if c in e['row'].index else '-'
        records.append(rec)

    return pd.DataFrame(records), all_time


# ── 그래프 생성 ───────────────────────────────────────────────────────────────

def split_title_subtitle(item_name: str):
    """'함량 (몬테루카스트 ...)' → ('함량', '(몬테루카스트 ...')')"""
    idx = item_name.find('(')
    if idx > 0:
        return item_name[:idx].strip(), item_name[idx:].strip()
    return item_name, ""


def _nice_impurity_ymax(candidate):
    for s in [0.05, 0.10, 0.15, 0.20, 0.30, 0.50, 1.0, 1.5, 2.0, 2.7, 3.0, 5.0]:
        if candidate <= s:
            return s
    return round(candidate * 10) / 10 + 0.1


def _is_special_value(text):
    """N.D 또는 부등호(<, >, ≤, ≥) 포함 값 여부 — 빗금 표시 기준.

    '-' 또는 빈 문자열(결측치)은 False를 반환해 일반 처리한다.
    """
    if not text or text in ('-', ''):
        return False
    s = text.strip().upper()
    return (s in {'N.D', 'ND', 'N.D.', 'NOT DETECTED',
                  '< D.L', '<D.L', '< D.L.', '< DL',
                  'DISREGARD LIMIT', '≤ DISREGARD LIMIT'}
            or s.startswith('<') or s.startswith('>')
            or '≤' in text or '≥' in text)


def create_3d_bar_chart(entries, chart_title, subtitle, limit_max=None):
    """
    유연물질 전용 3D 막대 그래프 + 하단 데이터 테이블 → PNG bytes.

    - N.D / < D.L: 연한 회색 막대 + Poly3DCollection 빗금(///) 오버레이
    - 수치 셀: 우측 정렬 / 특수 값 셀: 중앙 정렬 + 회색 배경
    - 배치명 열 너비 고정(20%), 나머지 균등 분배
    """
    from mpl_toolkits.mplot3d import Axes3D  # noqa — 3D projection 등록

    # ── Guard Clause ─────────────────────────────────────────────────────────
    if not entries:
        _log.warning("create_3d_bar_chart: entries가 비어 있어 차트를 생성하지 않습니다.")
        return None

    # entries 내부 유효성 확인
    valid_entries = [
        e for e in entries
        if isinstance(e, dict)
        and e.get('time_cols')
        and e.get('row') is not None
    ]
    if not valid_entries:
        _log.warning(
            "create_3d_bar_chart: 유효한 entry가 없습니다 (time_cols 또는 row 누락). "
            "entries=%r", entries
        )
        return None

    time_col_names = valid_entries[0]['time_cols']
    if not time_col_names:
        _log.warning("create_3d_bar_chart: time_col_names가 비어 있습니다.")
        return None

    n_time    = len(time_col_names)
    n_batches = len(valid_entries)

    # ── 데이터 수집 ──
    batch_data = []  # [(batch_no, [numeric], [display_str], [is_special]), ...]
    for e in valid_entries:
        nums, disps, specials = [], [], []
        for c in time_col_names:
            raw = str(e['row'].get(c, '')).strip()
            v = parse_value(raw)
            nums.append(v if v is not None else 0.0)
            disp = raw if raw and raw.lower() not in ('nan', '') else '-'
            disps.append(disp)
            specials.append(_is_special_value(disp))
        batch_data.append((e['batch'], nums, disps, specials))

    if not batch_data:
        _log.warning("create_3d_bar_chart: batch_data 수집 후 데이터가 없습니다.")
        return None

    # ── Z축 상한 ──
    all_nums = [v for _, vals, _, _ in batch_data for v in vals]
    data_max = max(all_nums) if all_nums else 0.0
    z_max    = _nice_impurity_ymax(max(data_max * 1.3, (limit_max or 0) * 1.2, 0.01))
    min_bar  = z_max * 0.012

    # ── Figure ──
    fig_h = 5.5 + (n_batches + 1) * 0.30
    fig   = plt.figure(figsize=(9 * 0.95, fig_h * 0.75), facecolor='white')

    try:
        fig.text(0.5, 0.97, chart_title, ha='center', va='top',
                 fontsize=14, fontweight='bold')
        fig.text(0.5, 0.92, subtitle, ha='center', va='top', fontsize=11)

        gs = plt.GridSpec(2, 1, figure=fig,
                          height_ratios=[5, 1],
                          top=0.88, bottom=0.04,
                          left=0.04, right=0.96,   # 좌우 여백 확보
                          hspace=0.08)
        ax   = fig.add_subplot(gs[0], projection='3d')
        ax_t = fig.add_subplot(gs[1])
        ax_t.axis('off')

        bw, bd = 0.68, 0.50   # 막대 두께 강화

        for bi, (batch_no, nums, disps, specials) in enumerate(batch_data):
            color = LINE_COLORS[bi % len(LINE_COLORS)]

            for ti, (v, d, is_special) in enumerate(zip(nums, disps, specials)):
                bh = max(v, min_bar)
                x0, y0 = float(ti), float(bi)

                if is_special:
                    # N.D / 부등호 값: 연한 회색 막대 + 크로스 해치
                    ax.bar3d(x0 - bw / 2, y0 - bd / 2, 0,
                             bw, bd, bh,
                             color='#d8d8d8', alpha=0.65, shade=True,
                             edgecolor='#1a1a4e', linewidth=1.5)

                    # 앞면 크로스 해치 (/ + \) — ax.plot3D() 직접 렌더링
                    y_f  = y0 - bd / 2
                    xmin = x0 - bw / 2
                    xmax = x0 + bw / 2
                    step = bw / 9           # 간격 좁힘 → 밀도↑
                    slope_up   =  bh / bw if bw > 0 else 1.0
                    slope_down = -bh / bw if bw > 0 else -1.0

                    for s_dir, s_val in ((slope_up, 1), (slope_down, -1)):
                        for k in range(-2, 11):
                            x_off = xmin + k * step
                            xa = max(xmin, x_off)
                            xb = min(xmax, x_off + bw)
                            if xb <= xa:
                                continue
                            if s_val == 1:
                                za = max(0.0, min(bh, s_dir * (xa - x_off)))
                                zb = max(0.0, min(bh, s_dir * (xb - x_off)))
                            else:
                                za = max(0.0, min(bh, bh + s_dir * (xa - x_off)))
                                zb = max(0.0, min(bh, bh + s_dir * (xb - x_off)))
                            ax.plot3D([xa, xb], [y_f, y_f], [za, zb],
                                      color='#1a1a4e', alpha=0.92, linewidth=2.0)
                else:
                    ax.bar3d(x0 - bw / 2, y0 - bd / 2, 0,
                             bw, bd, bh,
                             color=color, alpha=0.82, shade=True,
                             edgecolor='#444444', linewidth=0.8)

                label = d if d not in ('-', '') else ('N.D' if v == 0.0 else f'{v:.3f}%')
                ax.text(x0, y0, bh + z_max * 0.055,   # 5% 위쪽 오프셋
                        label, ha='center', va='bottom', fontsize=8, color='#222222')

        # 허용기준 반투명 평면 + 경계선
        if limit_max is not None and limit_max <= z_max:
            x_r = np.array([-0.5, n_time - 0.5])
            y_r = np.array([-0.5, max(n_batches - 0.5, 0.5)])
            xx, yy = np.meshgrid(x_r, y_r)
            zz = np.full_like(xx, limit_max, dtype=float)
            ax.plot_surface(xx, yy, zz, alpha=0.13, color='red', linewidth=0)
            # 경계선 — 4변을 명시적으로 그어 Upper Limit 위치 명확히 표시
            for xi in x_r:
                ax.plot3D([xi, xi], list(y_r), [limit_max, limit_max],
                          color='red', linewidth=1.6, alpha=0.75)
            for yi in y_r:
                ax.plot3D(list(x_r), [yi, yi], [limit_max, limit_max],
                          color='red', linewidth=1.6, alpha=0.75)

        # X축: 8pt, 45° 회전, 막대 중앙 정렬(ha='center')
        ax.set_xticks(range(n_time))
        ax.set_xticklabels(time_col_names, fontsize=8, rotation=45, ha='center')
        ax.set_yticks(range(n_batches))
        ax.set_yticklabels([d[0] for d in batch_data], fontsize=9)
        ax.set_zlim(0, z_max)
        ax.zaxis.set_major_formatter(ticker.FuncFormatter(lambda x, _: f'{x:.2f}%'))
        ax.tick_params(axis='z', labelsize=8)
        ax.set_xlabel('')
        ax.set_ylabel('')
        # azim 조정: 오른쪽 하단 여유 확보 (-42 → 우측 끝 라벨이 프레임 안쪽)
        ax.view_init(elev=28, azim=-42)

        # ── 하단 데이터 테이블 ──
        col_labels  = [''] + list(time_col_names)
        rows_data   = [[bn] + dv for bn, _, dv, _ in batch_data]
        batch_col_w = 0.20
        data_col_w  = (1.0 - batch_col_w) / n_time if n_time else 0.80
        col_widths  = [batch_col_w] + [data_col_w] * n_time

        tbl = ax_t.table(cellText=rows_data, colLabels=col_labels,
                         cellLoc='center', loc='center', bbox=[0, 0, 1, 1],
                         colWidths=col_widths)
        tbl.auto_set_font_size(False)
        tbl.set_fontsize(9)

        specials_grid = [[False] + list(sv) for _, _, _, sv in batch_data]

        for (r, c), cell in tbl.get_celld().items():
            cell.set_linewidth(0.5)
            cell.set_edgecolor('#aaaaaa')
            if r == 0:
                cell.set_facecolor('#dce6f1')
                cell.get_text().set_fontweight('bold')
                cell.get_text().set_fontsize(9)
                cell.get_text().set_ha('center')
            else:
                row_sp = specials_grid[r - 1] if r - 1 < len(specials_grid) else []
                is_sp  = row_sp[c] if c < len(row_sp) else False

                if c == 0:
                    cell.set_facecolor('white')
                    cell.get_text().set_color(LINE_COLORS[(r - 1) % len(LINE_COLORS)])
                    cell.get_text().set_fontweight('bold')
                    cell.get_text().set_ha('left')
                    cell.get_text().set_fontsize(9)
                elif is_sp:
                    cell.set_facecolor('#f0f0f0')
                    cell.get_text().set_ha('center')
                    cell.get_text().set_color('#888888')
                    cell.get_text().set_fontstyle('italic')
                    cell.get_text().set_fontsize(9)
                else:
                    cell.set_facecolor('white')
                    cell.get_text().set_ha('right')
                    cell.get_text().set_fontsize(9)

        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                    pad_inches=0.3,
                    facecolor='white', edgecolor='none')
        buf.seek(0)
        return buf.getvalue()

    except Exception as exc:
        _log.error(
            "create_3d_bar_chart: 차트 렌더링 중 오류 발생 — %s: %s",
            type(exc).__name__, exc, exc_info=True
        )
        return None

    finally:
        plt.close(fig)


def create_3d_ribbon_chart(entries, chart_title, subtitle, limit_max=None):
    """
    유연물질 전용 3D 리본 꺾은선 그래프 + 하단 테이블.

    각 배치를 Y 방향 두께를 가진 리본 면(Poly3DCollection)으로 표현하고
    위쪽 모서리에 선을 올려 예시 이미지 스타일을 재현.
    """
    from mpl_toolkits.mplot3d import Axes3D             # noqa — 3d projection 등록
    from mpl_toolkits.mplot3d.art3d import Poly3DCollection

    if not entries:
        return None

    time_col_names = entries[0]['time_cols']
    n_time     = len(time_col_names)
    n_batches  = len(entries)

    # ── 데이터 수집 ─────────────────────────────────────────────────────────
    batch_data = []          # [(batch_no, [numeric], [display_str]), ...]
    for e in entries:
        nums, disps = [], []
        for c in time_col_names:
            raw = str(e['row'].get(c, '')).strip()
            v   = parse_value(raw)
            nums.append(v if v is not None else 0.0)
            disps.append(raw if raw and raw.lower() not in ('nan', '') else '-')
        batch_data.append((e['batch'], nums, disps))

    # ── Z축 상한 (spec × 1.2) ───────────────────────────────────────────────
    all_nums = [v for _, vals, _ in batch_data for v in vals]
    data_max = max(all_nums) if all_nums else 0.0
    z_max = _nice_impurity_ymax(
        max(data_max * 1.3, (limit_max or 0) * 1.2, 0.01)
    )

    # ── Figure ──────────────────────────────────────────────────────────────
    fig_h = 5.5 + (n_batches + 1) * 0.28
    fig = plt.figure(figsize=(9 * 0.7, fig_h * 0.7), facecolor='white')

    fig.text(0.5, 0.97, chart_title, ha='center', va='top',
             fontsize=14, fontweight='bold')
    fig.text(0.5, 0.92, subtitle,    ha='center', va='top', fontsize=11)

    gs = plt.GridSpec(2, 1, figure=fig,
                      height_ratios=[5, 1],
                      top=0.88, bottom=0.04, hspace=0.08)
    ax   = fig.add_subplot(gs[0], projection='3d')
    ax_t = fig.add_subplot(gs[1])
    ax_t.axis('off')

    x_pos        = np.arange(n_time, dtype=float)
    ribbon_depth = 0.35    # 리본의 Y 방향 두께
    batch_gap    = 1.0     # 배치 간 Y 간격

    for bi, (batch_no, vals, disps) in enumerate(batch_data):
        color = LINE_COLORS[bi % len(LINE_COLORS)]
        y0 = bi * batch_gap            # 앞면 Y 좌표
        y1 = y0 + ribbon_depth         # 뒷면 Y 좌표

        # ── 리본을 구성하는 폴리곤 면들 ──
        verts = []
        for i in range(n_time - 1):
            xa, xb = x_pos[i], x_pos[i + 1]
            za, zb = vals[i], vals[i + 1]

            # 윗면 (데이터 수준의 수평 면)
            verts.append([(xa, y0, za), (xb, y0, zb),
                          (xb, y1, zb), (xa, y1, za)])
            # 앞면 (y=y0, 바닥~데이터)
            verts.append([(xa, y0, 0), (xb, y0, 0),
                          (xb, y0, zb), (xa, y0, za)])
            # 뒷면 (y=y1, 바닥~데이터)
            verts.append([(xa, y1, 0), (xb, y1, 0),
                          (xb, y1, zb), (xa, y1, za)])

        # 좌·우 끝 면 (end caps)
        verts.append([(x_pos[0],  y0, 0), (x_pos[0],  y1, 0),
                      (x_pos[0],  y1, vals[0]),  (x_pos[0],  y0, vals[0])])
        verts.append([(x_pos[-1], y0, 0), (x_pos[-1], y1, 0),
                      (x_pos[-1], y1, vals[-1]), (x_pos[-1], y0, vals[-1])])

        poly = Poly3DCollection(
            verts,
            facecolor=color, edgecolor=color,
            alpha=0.75, linewidth=0.4
        )
        ax.add_collection3d(poly)

        # 위쪽 모서리 선 (선명도 보강)
        ax.plot(x_pos, [y0] * n_time, vals, color=color, linewidth=2.0, zorder=5)
        ax.plot(x_pos, [y1] * n_time, vals, color=color, linewidth=1.0, zorder=5)

    # ── 허용기준 반투명 평면 ─────────────────────────────────────────────────
    if limit_max is not None and limit_max <= z_max:
        total_y = n_batches * batch_gap + ribbon_depth
        xr = np.array([-0.4, n_time - 0.6])
        yr = np.array([-0.1, total_y])
        xx, yy = np.meshgrid(xr, yr)
        zz = np.full_like(xx, limit_max, dtype=float)
        ax.plot_surface(xx, yy, zz, alpha=0.12, color='red', linewidth=0)

    # ── 축 설정 ─────────────────────────────────────────────────────────────
    ax.set_xticks(x_pos)
    ax.set_xticklabels(time_col_names, fontsize=10)

    y_centers = [bi * batch_gap + ribbon_depth / 2 for bi in range(n_batches)]
    ax.set_yticks(y_centers)
    ax.set_yticklabels([d[0] for d in batch_data], fontsize=9)

    ax.set_zlim(0, z_max)
    ax.zaxis.set_major_formatter(
        ticker.FuncFormatter(lambda x, _: f'{x:.1f}%'))
    ax.tick_params(axis='z', labelsize=9)
    ax.set_xlabel('')
    ax.set_ylabel('')

    # 패널을 흰 배경 + 연한 테두리로 (이미지 스타일)
    for pane in (ax.xaxis.pane, ax.yaxis.pane, ax.zaxis.pane):
        pane.fill = False
        pane.set_edgecolor('#cccccc')

    ax.view_init(elev=15, azim=-65)

    # ── 하단 데이터 테이블 ───────────────────────────────────────────────────
    col_labels = [''] + list(time_col_names)
    rows_data  = [[bn] + dv for bn, _, dv in batch_data]

    tbl = ax_t.table(cellText=rows_data, colLabels=col_labels,
                     cellLoc='center', loc='center', bbox=[0, 0, 1, 1])
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(10)
    for (r, c), cell in tbl.get_celld().items():
        cell.set_linewidth(0.5)
        cell.set_edgecolor('#aaaaaa')
        if r == 0:
            cell.set_facecolor('#dce6f1')
            cell.get_text().set_fontweight('bold')
        else:
            cell.set_facecolor('white')
            if c == 0:
                cell.get_text().set_color(LINE_COLORS[(r - 1) % len(LINE_COLORS)])
                cell.get_text().set_fontweight('bold')

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    buf.seek(0)
    img_bytes = buf.getvalue()
    plt.close(fig)
    return img_bytes
# ── end create_3d_ribbon_chart ────────────────────────────────────────────────


def create_chart(df, time_cols, chart_title, subtitle,
                 limit_min=None, limit_max=None, is_assay=True):
    """꺾은선 그래프 + 하단 데이터 테이블 → PNG bytes."""
    batch_col = df.columns[0]
    batches = [
        (str(row[batch_col]).strip(), row)
        for _, row in df.iterrows()
        if str(row[batch_col]).strip() and str(row[batch_col]).strip().lower() != 'nan'
    ]
    if not batches:
        return None

    all_vals = []
    for _, row in batches:
        for c in time_cols:
            v = parse_value(row.get(c))
            if v is not None:
                all_vals.append(v)
    if not all_vals:
        return None

    data_min, data_max = min(all_vals), max(all_vals)

    # Figure 크기: 배치 수에 따라 테이블 높이 조정
    n_batches = len(batches)
    table_rows = n_batches + 1
    fig_h = 5.5 + table_rows * 0.28
    fig = plt.figure(figsize=(9 * 0.7, fig_h * 0.7), facecolor='white')

    fig.text(0.5, 0.98, chart_title, ha='center', va='top',
             fontsize=14, fontweight='bold')
    fig.text(0.5, 0.93, subtitle, ha='center', va='top', fontsize=11)

    table_frac = min(0.30, 0.06 * table_rows)
    graph_bottom = 0.06 + table_frac
    ax = fig.add_axes([0.11, graph_bottom, 0.84, 0.88 - graph_bottom])

    x_pos = list(range(len(time_cols)))

    for idx, (batch_no, row) in enumerate(batches):
        xs, ys = [], []
        for i, c in enumerate(time_cols):
            v = parse_value(row.get(c))
            if v is not None:
                xs.append(x_pos[i])
                ys.append(v)
        if xs:
            color = LINE_COLORS[idx % len(LINE_COLORS)]
            ax.plot(xs, ys, 'o-', color=color, linewidth=2, markersize=6, zorder=3)

    if limit_min is not None:
        ax.axhline(y=limit_min, color='red', linestyle='--', linewidth=1.2, zorder=2)
    if limit_max is not None:
        ax.axhline(y=limit_max, color='red', linestyle='--', linewidth=1.2, zorder=2)

    # Y축 범위
    if is_assay:
        lo = np.floor(min(data_min, limit_min if limit_min else data_min) - 1.5)
        hi = np.ceil(max(data_max, limit_max if limit_max else data_max) + 1.5)
    else:
        lo = 0.0
        hi = _nice_impurity_ymax(max(data_max * 1.3, (limit_max or 0) * 1.15, 0.01))

    ax.set_ylim(lo, hi)
    ax.yaxis.set_major_formatter(ticker.FuncFormatter(lambda x, _: f'{x:.1f}%'))
    ax.set_xticks(x_pos)
    ax.set_xticklabels(time_cols, fontsize=10)
    ax.set_xlim(-0.3, len(time_cols) - 0.7)
    ax.tick_params(axis='y', labelsize=10)
    ax.grid(True, color='#cccccc', linewidth=0.7, linestyle='-', zorder=0)
    ax.set_axisbelow(True)
    for sp in ax.spines.values():
        sp.set_linewidth(0.8)
        sp.set_edgecolor('black')

    # 하단 데이터 테이블
    ax_t = fig.add_axes([0.11, 0.02, 0.84, table_frac - 0.02])
    ax_t.axis('off')

    col_labels = [''] + [str(c) for c in time_cols]
    rows_data = []
    for batch_no, row in batches:
        r = [batch_no]
        for c in time_cols:
            v = str(row.get(c, '-')).strip()
            r.append(v if v and v.lower() != 'nan' else '-')
        rows_data.append(r)

    # 배치명 열 너비 고정(20%), 나머지 균등 분배
    _bcw = 0.20
    _dcw = (1.0 - _bcw) / len(time_cols) if time_cols else 0.80
    tbl = ax_t.table(
        cellText=rows_data,
        colLabels=col_labels,
        cellLoc='center',
        loc='center',
        bbox=[0, 0, 1, 1],
        colWidths=[_bcw] + [_dcw] * len(time_cols),
    )
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(9)

    for (r, c), cell in tbl.get_celld().items():
        cell.set_linewidth(0.5)
        cell.set_edgecolor('#aaaaaa')
        if r == 0:
            cell.set_facecolor('#dce6f1')
            cell.get_text().set_fontweight('bold')
            cell.get_text().set_fontsize(9)
            cell.get_text().set_ha('center')
        else:
            if c == 0:
                cell.set_facecolor('white')
                cell.get_text().set_color(LINE_COLORS[(r - 1) % len(LINE_COLORS)])
                cell.get_text().set_fontweight('bold')
                cell.get_text().set_ha('left')
                cell.get_text().set_fontsize(9)
            else:
                txt = cell.get_text().get_text()
                if _is_special_value(txt):
                    cell.set_facecolor('#f0f0f0')
                    cell.get_text().set_ha('center')
                    cell.get_text().set_color('#888888')
                    cell.get_text().set_fontstyle('italic')
                    cell.get_text().set_fontsize(9)
                else:
                    cell.set_facecolor('white')
                    cell.get_text().set_ha('right')
                    cell.get_text().set_fontsize(9)

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    buf.seek(0)
    img_bytes = buf.getvalue()
    plt.close(fig)
    return img_bytes


# ── Streamlit UI ──────────────────────────────────────────────────────────────

CUSTOM_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@300;400;500;700&display=swap');

html, body, [class*="css"] {
    font-family: 'Noto Sans KR', 'Malgun Gothic', sans-serif !important;
}
.stApp {
    background-color: #f4f6f4 !important;
}

/* ── 헤더 카드 ── */
.app-header {
    background: linear-gradient(135deg, #006633 0%, #00843d 100%);
    color: white;
    padding: 1.8rem 2.4rem;
    border-radius: 14px;
    margin-bottom: 1.6rem;
    box-shadow: 0 6px 24px rgba(0, 102, 51, 0.22);
}
.app-header h1 { margin: 0; font-size: 1.7rem; font-weight: 700; letter-spacing: -0.5px; line-height: 1.0; }
.app-header p  { margin: 0; font-size: 0.88rem; opacity: 0.9; line-height: 1.6; }

/* ── 메트릭 카드 ── */
.metric-row { display: flex; gap: 1rem; margin: 1.2rem 0 1.5rem; }
.metric-card {
    flex: 1; background: white; border-radius: 10px;
    padding: 1rem 1.4rem;
    box-shadow: 0 2px 10px rgba(0,0,0,0.07);
    border-left: 4px solid #006633;
}
.metric-card .val { font-size: 2rem; font-weight: 700; color: #006633; line-height: 1.1; }
.metric-card .lbl { font-size: 0.82rem; color: #666; margin-top: 0.15rem; }

/* ── 섹션 헤더 ── */
.section-title {
    display: flex; align-items: center; gap: 0.6rem;
    background: linear-gradient(90deg, #006633 0%, #00843d 55%, #c8e6c9 100%);
    color: white; padding: 0.7rem 1.4rem; border-radius: 8px;
    margin: 2rem 0 1.2rem; font-size: 1.05rem; font-weight: 600;
}

/* ── 설정 패널 ── */
.settings-panel {
    background: #f9fdf9; border-radius: 10px;
    padding: 1.2rem 1.4rem; border: 1px solid #c8e6c9; margin-bottom: 1.2rem;
}

/* ── 차트 카드 ── */
.chart-card {
    background: white; border-radius: 12px;
    padding: 1.4rem; margin-bottom: 0.8rem;
    box-shadow: 0 3px 16px rgba(0,0,0,0.08); border: 1px solid #e0ede0;
}

/* ── 버튼 ── */
.stDownloadButton > button {
    background-color: #006633 !important; color: white !important;
    border: none !important; border-radius: 8px !important;
    padding: 0.45rem 1.4rem !important; font-weight: 500 !important;
    font-size: 0.9rem !important; transition: all 0.2s ease !important;
    box-shadow: 0 2px 8px rgba(0, 102, 51, 0.28) !important;
}
.stDownloadButton > button:hover {
    background-color: #004d26 !important;
    box-shadow: 0 4px 14px rgba(0, 102, 51, 0.38) !important;
    transform: translateY(-1px) !important;
}

/* ── 텍스트/숫자 입력 ── */
.stTextInput > div > div > input,
.stNumberInput > div > div > input {
    border-radius: 8px !important; border: 1.5px solid #c8e6c9 !important;
    background: white !important; transition: all 0.2s !important;
}
.stTextInput > div > div > input:focus,
.stNumberInput > div > div > input:focus {
    border-color: #006633 !important;
    box-shadow: 0 0 0 3px rgba(0, 102, 51, 0.13) !important;
}

/* ── 파일 업로더 ── */
[data-testid="stFileUploader"] {
    background: white; border-radius: 12px;
    border: 2px dashed #a5d6a7; padding: 0.5rem;
    transition: border-color 0.25s;
}
[data-testid="stFileUploader"]:hover { border-color: #006633; }

/* ── 익스팬더 ── */
.streamlit-expanderHeader {
    background: #f9fdf9 !important; border-radius: 8px !important;
    border: 1px solid #c8e6c9 !important; font-weight: 500 !important;
}

/* ── 성공·정보 메시지 ── */
div[data-testid="stNotification"] {
    border-radius: 8px !important;
}

/* ── 구분선 ── */
hr { border: none !important; border-top: 2px solid #e0ede0 !important; margin: 0.8rem 0 !important; }

/* ── 푸터 ── */
.app-footer {
    text-align: center; color: #aaa; font-size: 0.78rem;
    margin-top: 3rem; padding: 1.2rem;
    border-top: 1px solid #e0ede0;
}
</style>
"""

st.set_page_config(
    page_title="안정성 데이터 시각화 리포트 생성기",
    page_icon="💡",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

# ── 헤더 ──────────────────────────────────────────────────────────────────────
st.markdown(
'<div style="background:linear-gradient(100deg,#004d26 0%,#008037 100%);padding:40px 20px;border-radius:15px;text-align:center;box-shadow:0 4px 15px rgba(0,0,0,0.1);color:white;">'
'<h1 style="margin:0;font-family:\'Segoe UI\',Roboto,Helvetica,Arial,sans-serif;font-size:2.8rem;font-weight:800;letter-spacing:-0.5px;">Stability <span style="font-weight:300;opacity:0.9;">Graph Generator</span></h1>'
'<div style="margin-top:10px;font-size:0.85rem;letter-spacing:2px;font-weight:400;opacity:0.7;">QUALITY CONTROL INTELLIGENCE SYSTEM</div>'
'</div><br>',
unsafe_allow_html=True)

# ── 파일 업로드 ───────────────────────────────────────────────────────────────
uploaded_file = st.file_uploader(
    "📂  워드 파일(.docx)을 선택하거나 드래그하세요",
    type=["docx"],
)

if uploaded_file:
    try:
        assay_rows, imp_groups, imp_limits, diag, extracted_batch = process_document(uploaded_file)

        # ── 메트릭 카드 ─────────────────────────────────────────────────────
        st.markdown(f"""
        <div class="metric-row">
          <div class="metric-card"><div class="val">{diag['n_tables']}</div><div class="lbl">📋 발견된 표 수</div></div>
          <div class="metric-card"><div class="val">{len(assay_rows)}</div><div class="lbl">💊 함량 데이터 행</div></div>
          <div class="metric-card"><div class="val">{len(imp_groups)}</div><div class="lbl">🔬 유연물질 항목 수</div></div>
        </div>
        """, unsafe_allow_html=True)

        if extracted_batch:
            st.markdown(f"""
            <div class="section-title">
              🏷️ 대상 배치 번호 : {extracted_batch}
            </div>
            """, unsafe_allow_html=True)

        # ── 표 구조 진단 ─────────────────────────────────────────────────────
        with st.expander("🔍 표 구조 진단 상세 보기",
                         expanded=not (assay_rows or imp_groups)):
            st.caption("파싱 대상: doc.tables[1] (두 번째 표부터)")
            if diag['raw_header']:
                st.write("**헤더 행 (index 1):**")
                st.dataframe(
                    pd.DataFrame([diag['raw_header']],
                                 columns=[f"col{i}" for i in range(len(diag['raw_header']))]),
                    use_container_width=True,
                )
            if diag['raw_rows']:
                st.write("**데이터 행 (앞 5행):**")
                st.dataframe(
                    pd.DataFrame(
                        diag['raw_rows'],
                        columns=[f"col{i}" for i in range(len(diag['raw_rows'][0]))]
                        if diag['raw_rows'] else [],
                    ),
                    use_container_width=True,
                )
            st.caption(
                "col1=대분류(index 1) · col2=세부항목(index 2) · "
                "col3=허용기준(index 3) · col4+=시점 데이터"
            )

        if not assay_rows and not imp_groups:
            st.error(
                "⚠️ 함량 또는 유연물질 데이터를 찾을 수 없습니다.  \n"
                "**표 구조 진단**에서 col2(index 2)에 "
                "'함량'·'유연물질'·'불순물' 키워드가 있는지 확인하세요."
            )
            st.stop()

        # ── 함량 그래프 ──────────────────────────────────────────────────────
        if assay_rows:
            st.markdown('<div class="section-title">💊 함량 (Assay)</div>',
                        unsafe_allow_html=True)

            df_assay, time_cols = build_dataframe(assay_rows)

            lmin_auto, lmax_auto = None, None
            for e in assay_rows:
                if e['limit_text']:
                    a, b = parse_limit_from_text(e['limit_text'])
                    if a is not None: lmin_auto = a
                    if b is not None: lmax_auto = b

            full_name = assay_rows[0]['item_name'] if assay_rows else "함량"
            t1, t2 = split_title_subtitle(full_name)
            if not t2:
                t2 = f"(표시량의 {lmin_auto or 94.0:.1f} %~{lmax_auto or 105.0:.1f} %)"

            st.markdown('<div class="settings-panel">', unsafe_allow_html=True)
            c1, c2 = st.columns(2)
            with c1:
                title_inp = st.text_input("📌 그래프 제목", value=t1, key="a_title")
            with c2:
                sub_inp = st.text_input("📋 부제목 (허용기준)", value=t2, key="a_sub")
            c3, c4, _ = st.columns([1, 1, 2])
            with c3:
                lmin = st.number_input("하한 (%)", value=float(lmin_auto or 94.0),
                                       step=0.1, format="%.1f", key="a_lmin")
            with c4:
                lmax = st.number_input("상한 (%)", value=float(lmax_auto or 105.0),
                                       step=0.1, format="%.1f", key="a_lmax")
            st.markdown('</div>', unsafe_allow_html=True)

            with st.expander("📊 원본 데이터 확인"):
                st.dataframe(df_assay, use_container_width=True)

            img = create_chart(df_assay, time_cols, title_inp, sub_inp,
                               limit_min=lmin, limit_max=lmax, is_assay=True)
            if img:
                st.markdown('<div class="chart-card">', unsafe_allow_html=True)
                col_img, _ = st.columns([0.6, 0.4])
                with col_img:
                    st.image(img, use_container_width=True)
                st.markdown('</div>', unsafe_allow_html=True)
                st.download_button("📥 PNG 다운로드 (함량)", data=img,
                                   file_name="함량_그래프.png", mime="image/png",
                                   key="dl_assay")

        # ── 유연물질 그래프 ───────────────────────────────────────────────────
        for i, (imp_name, entries) in enumerate(imp_groups.items()):
            st.markdown(
                f'<div class="section-title">🔬 유연물질 — {imp_name}</div>',
                unsafe_allow_html=True,
            )

            df_imp, _ = build_dataframe(entries)
            limit_text = imp_limits.get(imp_name, "")
            _, lmax_auto = parse_limit_from_text(limit_text) if limit_text else (None, None)

            t1, t2 = split_title_subtitle(imp_name)
            if not t2:
                t2 = f"(≤ {lmax_auto:.2f} %)" if lmax_auto else f"({imp_name})"

            st.markdown('<div class="settings-panel">', unsafe_allow_html=True)
            c1, c2 = st.columns(2)
            with c1:
                title_inp = st.text_input("📌 그래프 제목", value=t1,
                                          key=f"imp_title_{i}")
            with c2:
                sub_inp = st.text_input("📋 부제목 (허용기준)", value=t2,
                                        key=f"imp_sub_{i}")
            c3, _ = st.columns([1, 3])
            with c3:
                lmax = st.number_input("상한 (%)", value=float(lmax_auto or 2.7),
                                       step=0.01, format="%.2f",
                                       key=f"imp_lmax_{i}")
            st.markdown('</div>', unsafe_allow_html=True)

            with st.expander("📊 원본 데이터 확인"):
                st.dataframe(df_imp, use_container_width=True)

            img = create_3d_bar_chart(entries, title_inp, sub_inp, limit_max=lmax)
            if img:
                st.markdown('<div class="chart-card">', unsafe_allow_html=True)
                col_img, _ = st.columns([0.6, 0.4])
                with col_img:
                    st.image(img, use_container_width=True)
                st.markdown('</div>', unsafe_allow_html=True)
                st.download_button(
                    f"📥 PNG 다운로드 (유연물질 {i + 1})", data=img,
                    file_name=f"유연물질_{imp_name.replace(' ', '_')[:30]}.png",
                    mime="image/png",
                    key=f"dl_imp_{i}",
                )

    except Exception as exc:
        st.error(f"⚠️ 오류 발생: {exc}")
        import traceback
        st.code(traceback.format_exc())

# ── 푸터 ──────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="app-footer">
  안정성 그래프 생성기 · Stability Graph Generator · Powered by Streamlit &amp; Matplotlib
</div>
""", unsafe_allow_html=True)
